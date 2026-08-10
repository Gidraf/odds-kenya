import io
import time
from flask import request
from .blueprint import bp_odds_customer
from .utils import _apply_tier_limits, _now_utc, _normalise_sport_slug, _bk_slug, _flatten_db_markets, _effective_status
from .formatters import _build_envelope, _build_analytics_full, _build_analytics_summary
from .db_services import _load_db_matches, _get_analytics, _build_base_query, _sport_filter, _mode_time_filter, _multi_bk_filter
from .cache_services import _read_cache_sources, _deduplicate, _normalise_cache_match

@bp_odds_customer.route("/odds/upcoming/<sport_slug>")
def get_upcoming(sport_slug: str):
    from app.utils.customer_jwt_helpers import _current_user_from_header, _signed_response
    from app.utils.decorators_ import log_event
    t0 = time.perf_counter()
    user = _current_user_from_header()
    tier = getattr(user, "tier", "free") if user else "free"
    page = max(1, int(request.args.get("page", 1)))
    per_page = min(100, int(request.args.get("per_page", 20)))
    
    log_event("odds_upcoming", {"sport": sport_slug, "tier": tier})

    try:
        matches, total, pages = _load_db_matches(
            sport_slug, mode="upcoming", page=page, per_page=per_page,
            comp_filter=(request.args.get("comp", "") or "").strip(), team_filter=(request.args.get("team", "") or "").strip(),
            has_arb=request.args.get("has_arb", "") in ("1", "true"), sort=request.args.get("sort", "start_time"),
            date_str=request.args.get("date", ""), from_dt=request.args.get("from_dt", ""), to_dt=request.args.get("to_dt", ""),
            include_analytics=request.args.get("analytics","") in ("1", "true"),
        )
    except Exception:
        matches = [x for m in _deduplicate(_read_cache_sources("upcoming", sport_slug)) if (x := _normalise_cache_match(m, "upcoming")) is not None]
        total = len(matches); pages = max(1, (total + per_page - 1) // per_page)
        matches = matches[(page-1)*per_page: page*per_page]

    matches, truncated = _apply_tier_limits(matches, user)
    return _signed_response(_build_envelope(matches, sport_slug, "upcoming", tier, page, per_page, truncated, int((time.perf_counter() - t0) * 1000), total=total, pages=pages), encrypt_for=user)

@bp_odds_customer.route("/odds/_customer/live/<sport_slug>")
def get_live(sport_slug: str):
    from app.utils.customer_jwt_helpers import _current_user_from_header, _signed_response
    from app.utils.decorators_ import log_event
    t0 = time.perf_counter()
    user = _current_user_from_header()
    tier = getattr(user, "tier", "free") if user else "free"
    page = max(1, int(request.args.get("page", 1)))
    per_page = min(100, int(request.args.get("per_page", 20)))
    
    log_event("odds_live", {"sport": sport_slug, "tier": tier})

    try:
        matches, total, pages = _load_db_matches(sport_slug, mode="live", page=page, per_page=per_page, comp_filter=(request.args.get("comp", "") or "").strip(), team_filter=(request.args.get("team", "") or "").strip(), sort=request.args.get("sort", "start_time"), include_analytics=request.args.get("analytics","") in ("1", "true"))
    except Exception:
        matches = [x for m in _deduplicate(_read_cache_sources("live", sport_slug)) if (x := _normalise_cache_match(m, "live")) is not None]
        total = len(matches); pages = max(1, (total + per_page - 1) // per_page)
        matches = matches[(page-1)*per_page: page*per_page]

    matches, truncated = _apply_tier_limits(matches, user)
    return _signed_response(_build_envelope(matches, sport_slug, "live", tier, page, per_page, truncated, int((time.perf_counter() - t0) * 1000), total=total, pages=pages), encrypt_for=user)

@bp_odds_customer.route("/odds/results")
@bp_odds_customer.route("/odds/results/<date_str>")
def get_results_by_date(date_str: str = ""):
    from app.utils.customer_jwt_helpers import _current_user_from_header, _signed_response
    from app.utils.decorators_ import log_event
    from app.workers.celery_tasks import cache_get
    if not date_str: date_str = _now_utc().strftime("%Y-%m-%d")
    t0 = time.perf_counter()
    user = _current_user_from_header()
    tier = getattr(user, "tier", "free") if user else "free"
    page = max(1, int(request.args.get("page", 1)))
    per_page = min(100, int(request.args.get("per_page", 20)))
    
    log_event("finished_games_view", {"date": date_str})
    try: matches, total, pages = _load_db_matches(request.args.get("sport", "") or "all", mode="finished", page=page, per_page=per_page, date_str=date_str, comp_filter=(request.args.get("competition") or ""), team_filter=(request.args.get("team") or ""))
    except Exception:
        matches = [x for m in (cache_get(f"results:finished:{date_str}") or []) if (x := _normalise_cache_match(m, "finished")) is not None]
        total = len(matches); pages = max(1, (total + per_page - 1) // per_page)
        matches = matches[(page-1)*per_page: page*per_page]
        
    matches, truncated = _apply_tier_limits(matches, user)
    return _signed_response(_build_envelope(matches, date_str, "finished", tier, page, per_page, truncated, int((time.perf_counter() - t0) * 1000), total=total, pages=pages, extra={"date": date_str}), encrypt_for=user)

@bp_odds_customer.route("/odds/match/<parent_match_id>")
def get_match(parent_match_id: str):
    from app.utils.customer_jwt_helpers import _current_user_from_header, _err, _signed_response
    from app.utils.decorators_ import log_event
    from app.models.odds import UnifiedMatch, BookmakerMatchOdds, ArbitrageOpportunity, EVOpportunity
    from app.models.bookmakers_model import Bookmaker, BookmakerMatchLink
    from sqlalchemy import and_, func
    from datetime import timedelta, timezone

    t0 = time.perf_counter()
    user = _current_user_from_header()
    tier = getattr(user, "tier", "free") if user else "free"

    um = UnifiedMatch.query.filter_by(parent_match_id=parent_match_id).first()
    if not um: return _err("Match not found", 404)
    log_event("match_view", {"match_id": parent_match_id, "tier": tier})

    class _SiblingBMO:
        def __init__(self, bmo): self._bmo = bmo
        def __getattr__(self, name): return getattr(self._bmo, name)

    bmos = list(BookmakerMatchOdds.query.filter_by(match_id=um.id).all())
    try:
        if um.home_team_name and um.away_team_name:
            conds = [func.lower(UnifiedMatch.home_team_name) == um.home_team_name.lower().strip(), func.lower(UnifiedMatch.away_team_name) == um.away_team_name.lower().strip(), UnifiedMatch.id != um.id]
            if um.start_time: conds += [UnifiedMatch.start_time >= um.start_time - timedelta(minutes=90), UnifiedMatch.start_time <= um.start_time + timedelta(minutes=90)]
            for sib in UnifiedMatch.query.filter(and_(*conds)).all():
                for b in BookmakerMatchOdds.query.filter_by(match_id=sib.id).all(): bmos.append(_SiblingBMO(b))
    except Exception: pass

    bk_ids = {bmo.bookmaker_id for bmo in bmos}
    bk_map = ({b.id: b for b in Bookmaker.query.filter(Bookmaker.id.in_(bk_ids)).all()} if bk_ids else {})
    links = {lnk.bookmaker_id: lnk.to_dict() for lnk in BookmakerMatchLink.query.filter_by(match_id=um.id).all()}

    bookmakers = {}; markets_by_bk = {}
    for bmo in bmos:
        bk_obj = bk_map.get(bmo.bookmaker_id)
        sl = _bk_slug((bk_obj.name if bk_obj else str(bmo.bookmaker_id)).lower())
        mkts = _flatten_db_markets(bmo.markets_json or {})
        if not mkts: continue
        if sl in bookmakers:
            bookmakers[sl]["markets"].update(mkts); markets_by_bk[sl].update(mkts)
        else:
            bookmakers[sl] = {"bookmaker_id": bmo.bookmaker_id, "bookmaker": bk_obj.name if bk_obj else sl.upper(), "slug": sl, "markets": mkts, "market_count": len(mkts), "link": links.get(bmo.bookmaker_id)}
            markets_by_bk[sl] = mkts

    best = {}
    for sl, bk_mkts in markets_by_bk.items():
        for mkt, outcomes in bk_mkts.items():
            best.setdefault(mkt, {})
            for out, odd_data in (outcomes or {}).items():
                try: fv = (float(odd_data.get("price") or odd_data.get("odd") or 0) if isinstance(odd_data, dict) else float(odd_data))
                except Exception: continue
                if fv > 1.0 and (out not in best[mkt] or fv > best[mkt][out]["odd"]): best[mkt][out] = {"odd": fv, "bk": sl}

    status_out = _effective_status(getattr(um, "status", None), um.start_time)
    minutes_elapsed = int((_now_utc() - (um.start_time if um.start_time.tzinfo else um.start_time.replace(tzinfo=timezone.utc))).total_seconds() / 60) if um.start_time and status_out == "IN_PLAY" else None

    history_rows = []
    
    try:
        arb_list = [a.to_dict() for a in ArbitrageOpportunity.query.filter_by(match_id=um.id, status="OPEN").all()]
        ev_list  = [e.to_dict() for e in EVOpportunity.query.filter_by(match_id=um.id, status="OPEN").all()]
    except Exception: arb_list = ev_list = []

    br_id = um.parent_match_id or ""
    analytics_bundle  = _get_analytics(br_id, trigger_if_missing=True)

    return _signed_response({
        "ok": True, "match_id": um.id, "parent_match_id": br_id, "betradar_id": br_id, "join_key": f"br_{br_id}" if br_id else f"db_{um.id}",
        "home_team": um.home_team_name, "away_team": um.away_team_name, "competition": um.competition_name, "sport": _normalise_sport_slug(um.sport_name or ""),
        "start_time": um.start_time.isoformat() if um.start_time else None, "status": status_out, "is_live": status_out == "IN_PLAY", "minutes_elapsed": minutes_elapsed,
        "bookmakers": bookmakers, "markets_by_bk": markets_by_bk, "markets": markets_by_bk, "best": best, "aggregated": _flatten_db_markets(um.markets_json or {}),
        "odds_history": history_rows, "arbs": arb_list, "evs": ev_list, "bk_ids": {sl: str(d["bookmaker_id"]) for sl, d in bookmakers.items()},
        "has_analytics": _build_analytics_full(analytics_bundle).get("available", False), "analytics": _build_analytics_full(analytics_bundle), "analytics_summary": _build_analytics_summary(analytics_bundle),
        "latency_ms": int((time.perf_counter() - t0) * 1000), "server_time": _now_utc().isoformat(), "source": "postgresql",
    }, encrypt_for=user)

@bp_odds_customer.route("/odds/match/<parent_match_id>/markets")
def get_match_full_markets(parent_match_id: str):
    from app.utils.customer_jwt_helpers import _current_user_from_header, _err, _signed_response
    from app.models.odds import UnifiedMatch, BookmakerMatchOdds
    from app.models.bookmakers_model import Bookmaker
    t0 = time.perf_counter()
    user = _current_user_from_header()

    um = UnifiedMatch.query.filter_by(parent_match_id=parent_match_id).first()
    if not um: return _err("Match not found", 404)

    bmo_rows = BookmakerMatchOdds.query.filter_by(match_id=um.id).all()
    all_bk_ids = {bmo.bookmaker_id for bmo in bmo_rows}
    bk_map = ({b.id: b for b in Bookmaker.query.filter(Bookmaker.id.in_(all_bk_ids)).all()} if all_bk_ids else {})

    stored_markets = {}
    for bmo in bmo_rows:
        if bk_obj := bk_map.get(bmo.bookmaker_id): stored_markets[_bk_slug(bk_obj.name.lower())] = _flatten_db_markets(bmo.markets_json or {})

    fresh_bt = {}; fetch_errors = {}
    try:
        from app.workers.bt_harvester import get_full_markets
        fresh_bt = get_full_markets(parent_match_id, sport_slug=_normalise_sport_slug(um.sport_name or "soccer"))
    except Exception as exc: fetch_errors["bt"] = str(exc)

    combined = dict(stored_markets)
    if fresh_bt: combined["bt"] = {**(stored_markets.get("bt") or {}), **fresh_bt}

    best = {}
    for sl, mkts in combined.items():
        for mkt, outcomes in (mkts or {}).items():
            best.setdefault(mkt, {})
            for out, odd_data in (outcomes or {}).items():
                try: fv = (float(odd_data.get("price") or odd_data.get("odd") or 0) if isinstance(odd_data, dict) else float(odd_data))
                except Exception: continue
                if fv > 1.0 and (out not in best[mkt] or fv > best[mkt][out]["odd"]): best[mkt][out] = {"odd": fv, "bk": sl}

    arb_markets = []
    if len(combined) >= 2:
        try:
            from app.workers.arb_engine import detect_arb_for_stream
            has_arb, _best_arb_pct, arb_markets = detect_arb_for_stream(best)
            if not has_arb:
                arb_markets = []
            else:
                for a in arb_markets:
                    if "market_slug" not in a:
                        a["market_slug"] = a.get("market")
                    if "market_label" not in a:
                        a["market_label"] = a.get("market_display") or a.get("market")
        except Exception:
            arb_markets = []

    status_out = _effective_status(getattr(um, "status", None), um.start_time)
    analytics_summary = _build_analytics_summary(_get_analytics(parent_match_id, trigger_if_missing=True))

    return _signed_response({
        "ok": True, "match_id": um.id, "parent_match_id": parent_match_id, "home_team": um.home_team_name, "away_team": um.away_team_name, "competition": um.competition_name,
        "sport": _normalise_sport_slug(um.sport_name or ""), "start_time": um.start_time.isoformat() if um.start_time else None, "status": status_out, "is_live": status_out == "IN_PLAY",
        "markets_by_bk": combined, "best": best, "arb_markets": arb_markets, "has_arb": bool(arb_markets), "market_count": len(best), "market_slugs": sorted(best.keys()),
        "bk_market_counts": {sl: len(mkts) for sl, mkts in combined.items()}, "bookmakers": sorted(combined.keys()), "bt_live_fetched": bool(fresh_bt), "bt_external_id": parent_match_id,
        "stored_bookmakers": sorted(stored_markets.keys()), "fetch_errors": fetch_errors, "has_analytics": analytics_summary.get("available", False), "analytics": analytics_summary,
        "latency_ms": int((time.perf_counter() - t0) * 1000), "server_time": _now_utc().isoformat(),
    }, encrypt_for=user)

@bp_odds_customer.route("/odds/match/<parent_match_id>/analytics")
def get_match_analytics(parent_match_id: str):
    from app.utils.customer_jwt_helpers import _current_user_from_header, _err, _signed_response
    from app.utils.decorators_ import log_event
    from app.models.odds import UnifiedMatch
    t0 = time.perf_counter()
    user = _current_user_from_header()
    um = UnifiedMatch.query.filter_by(parent_match_id=parent_match_id).first()
    if not um: return _err("Match not found", 404)
    log_event("match_analytics_view", {"match_id": parent_match_id})

    analytics = _build_analytics_full(_get_analytics(parent_match_id, trigger_if_missing=True))
    return _signed_response({
        "ok": True, "match_id": um.id, "parent_match_id": parent_match_id, "betradar_id": parent_match_id, "home_team": um.home_team_name, "away_team": um.away_team_name,
        "competition": um.competition_name, "sport": _normalise_sport_slug(um.sport_name or ""), "start_time": um.start_time.isoformat() if um.start_time else None,
        "available": analytics.get("available", False), "fetching": not analytics.get("available", False), "analytics": analytics, "latency_ms": int((time.perf_counter() - t0) * 1000), "server_time": _now_utc().isoformat(),
    }, encrypt_for=user)

@bp_odds_customer.route("/odds/match/<parent_match_id>/analytics/refresh", methods=["POST"])
def refresh_match_analytics(parent_match_id: str):
    from app.utils.customer_jwt_helpers import _err, _signed_response
    from app.models.odds import UnifiedMatch
    from app.workers.celery_tasks import celery
    if not UnifiedMatch.query.filter_by(parent_match_id=parent_match_id).first(): return _err("Match not found", 404)
    try: celery.send_task("tasks.sp.get_match_analytics", args=[parent_match_id, True], queue="harvest", countdown=0); dispatched = True
    except Exception: dispatched = False
    return _signed_response({"ok": dispatched, "parent_match_id": parent_match_id, "dispatched": dispatched, "message": "Analytics refresh queued.", "server_time": _now_utc().isoformat()})

@bp_odds_customer.route("/odds/search")
def search_matches():
    from app.utils.customer_jwt_helpers import _current_user_from_header, _err, _signed_response
    from app.utils.decorators_ import log_event
    from app.models.odds import UnifiedMatch, BookmakerMatchOdds
    from sqlalchemy import or_, func as sqlfunc
    from app.extensions import db
    t0 = time.perf_counter()
    user = _current_user_from_header()
    tier = getattr(user, "tier", "free") if user else "free"
    q_str = (request.args.get("q") or "").strip()
    mode = request.args.get("mode", "upcoming")
    page = max(1, int(request.args.get("page", 1)))
    per_page = min(100, int(request.args.get("per_page", 20)))
    sport = (request.args.get("sport") or "").strip()
    if not q_str: return _err("Provide query param 'q'", 400)

    qs = UnifiedMatch.query.filter(or_(UnifiedMatch.home_team_name.ilike(f"%{q_str}%"), UnifiedMatch.away_team_name.ilike(f"%{q_str}%"), UnifiedMatch.competition_name.ilike(f"%{q_str}%"), UnifiedMatch.parent_match_id.ilike(f"%{q_str}%")))
    if sport: qs = _sport_filter(qs, sport)
    qs = _mode_time_filter(qs, mode)
    if mode in ("upcoming", "live"): qs = _multi_bk_filter(qs)
    
    total = qs.count()
    um_list = qs.order_by(UnifiedMatch.start_time).offset((page-1)*per_page).limit(per_page).all()
    match_ids = [um.id for um in um_list]
    bk_counts = dict(db.session.query(BookmakerMatchOdds.match_id, sqlfunc.count(BookmakerMatchOdds.bookmaker_id)).filter(BookmakerMatchOdds.match_id.in_(match_ids)).group_by(BookmakerMatchOdds.match_id).all()) if match_ids else {}

    results = [{"match_id": um.id, "parent_match_id": um.parent_match_id, "betradar_id": um.parent_match_id, "join_key": f"br_{um.parent_match_id}" if um.parent_match_id else f"db_{um.id}", "home_team": um.home_team_name, "away_team": um.away_team_name, "competition": um.competition_name, "sport": _normalise_sport_slug(um.sport_name or ""), "start_time": um.start_time.isoformat() if um.start_time else None, "status": _effective_status(getattr(um, "status", None), um.start_time), "is_live": _is_live(getattr(um, "status", None), um.start_time), "bookie_count": bk_counts.get(um.id, 0), "detail_url": f"/api/odds/match/{um.parent_match_id}", "analytics_url": f"/api/odds/match/{um.parent_match_id}/analytics"} for um in um_list]
    log_event("odds_search", {"q": q_str, "mode": mode, "total": total})
    return _signed_response({"ok": True, "q": q_str, "mode": mode, "tier": tier, "total": total, "page": page, "per_page": per_page, "pages": max(1, (total + per_page - 1) // per_page), "latency_ms": int((time.perf_counter() - t0) * 1000), "matches": results, "source": "postgresql"}, encrypt_for=user)

# ── Bookmaker slug → full display name ──────────────────────────────────────
_BK_DISPLAY = {
    "sp":        "SportPesa",
    "bt":        "Betika",
    "mz":        "Mozzart",
    "od":        "OdiBets",
    "1xbet":     "1xBet",
    "22bet":     "22Bet",
    "betwinner": "BetWinner",
    "melbet":    "Melbet",
    "megapari":  "Megapari",
    "helabet":   "Helabet",
    "paripesa":  "PariPesa",
    "sbo":       "SBO",
}

def _bk_display(slug: str) -> str:
    """Return the full bookmaker display name for a given slug."""
    return _BK_DISPLAY.get(slug.lower(), slug.upper())


# ── MinIO helpers for pre-generated Word document caching ───────────────────
def _get_minio_client():
    """Return an initialised MinIO client, or None if unavailable."""
    import os
    try:
        from minio import Minio
        endpoint = os.environ.get("STORAGE_ENDPOINT", "5.78.137.59:6500")
        clean    = endpoint.replace("http://", "").replace("https://", "")
        secure   = endpoint.startswith("https")
        client   = Minio(
            clean,
            access_key=os.environ.get("STORAGE_ACCESS_KEY", "minioadmin"),
            secret_key=os.environ.get("STORAGE_SECRET_KEY", "minioadmin"),
            secure=secure,
        )
        bucket = "odds-reports"
        if not client.bucket_exists(bucket):
            client.make_bucket(bucket)
        return client, bucket
    except Exception:
        return None, None


def _minio_report_key(sport: str, arb_only: bool, preset: str = "all") -> str:
    suffix = "_arb" if arb_only else "_full"
    return f"reports/{sport}_{preset}{suffix}_latest.docx"


def _serve_minio_report(sport: str, arb_only: bool, preset: str = "all"):
    """Try to fetch the pre-generated report from MinIO. Returns BytesIO or None."""
    try:
        client, bucket = _get_minio_client()
        if not client:
            return None
        key  = _minio_report_key(sport, arb_only, preset=preset)
        resp = client.get_object(bucket, key)
        data = resp.read()
        resp.close(); resp.release_conn()
        return io.BytesIO(data)
    except Exception:
        return None


def _save_minio_report(sport: str, arb_only: bool, buf: io.BytesIO, preset: str = "all") -> bool:
    """Upload a generated Word report to MinIO. buf position is preserved."""
    try:
        client, bucket = _get_minio_client()
        if not client:
            return False
        key    = _minio_report_key(sport, arb_only, preset=preset)
        buf.seek(0)
        length = buf.getbuffer().nbytes
        client.put_object(
            bucket, key, buf, length,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        buf.seek(0)
        return True
    except Exception:
        return False
def _generate_word_document(sport: str, arb_only: bool, start_time_str: str = None, end_time_str: str = None) -> io.BytesIO:  # noqa: C901
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    import io, time, json as _json
    from datetime import datetime as _dt, timezone as _tz, timedelta as _td

    # ── Time Parsing ──────────────────────────────────────────────────────────
    def _parse_dt(s):
        if not s: return None
        try: return _dt.fromisoformat(str(s).strip().replace("Z", "+00:00"))
        except Exception: return None

    start_dt = _parse_dt(start_time_str)
    end_dt   = _parse_dt(end_time_str)
    _now_dt  = _dt.now(_tz.utc)
    if start_dt is None:        start_dt = _now_dt
    elif start_dt < _now_dt:   start_dt = _now_dt

    # ── Style Constants ───────────────────────────────────────────────────────
    FF        = "Arial"
    PT_TITLE  = Pt(14)
    PT_SUB    = Pt(7.5)
    PT_SEC    = Pt(8.5)
    PT_BAND   = Pt(7)
    PT_TEXT   = Pt(7.5)
    PT_TINY   = Pt(6.5)
    PT_HDR    = Pt(7)

    RGB_W     = RGBColor(0xFF, 0xFF, 0xFF)
    RGB_TXT   = RGBColor(0x1E, 0x29, 0x3B)
    RGB_MUT   = RGBColor(0x64, 0x74, 0x8B)
    RGB_GRN   = RGBColor(0x15, 0x80, 0x3D)
    RGB_SP    = RGBColor(0x25, 0x63, 0xEB)
    RGB_BT    = RGBColor(0x16, 0xA3, 0x4A)
    RGB_OD    = RGBColor(0xD9, 0x77, 0x06)
    RGB_SKY   = RGBColor(0x7D, 0xD3, 0xFC)
    RGB_AMB   = RGBColor(0xFB, 0xBF, 0x24)
    RGB_EMD   = RGBColor(0x6E, 0xE7, 0xB7)
    RGB_VIO   = RGBColor(0xC4, 0xB5, 0xFD)
    RGB_ARB   = RGBColor(0xA7, 0xF3, 0xD0)

    HEX_PRI    = "0F172A"   # Deep navy
    HEX_ALT    = "F8FAFC"   # Light gray rows
    HEX_BEST   = "DCFCE7"   # Best-odd green tint
    HEX_BORDER = "CBD5E1"   # Subtle border

    # Time bands: (start_hour_EAT, end_hour_EAT, label, hex_bg, rgb_label_color)
    EAT = _td(hours=3)
    BANDS = [
        ( 0,  6, "🌙  LATE NIGHT",     "051B35", RGB_SKY),
        ( 6, 10, "🌅  EARLY MORNING",  "062318", RGB_EMD),
        (10, 14, "☀️   MORNING",        "1C1100", RGB_AMB),
        (14, 18, "🌤  AFTERNOON",       "1A0038", RGB_VIO),
        (18, 21, "🌆  EVENING",         "0C1A35", RGB_SKY),
        (21, 24, "🌙  LATE NIGHT",     "051B35", RGB_SKY),
    ]

    def _band_info(dt_utc):
        try:
            if dt_utc.tzinfo is None: dt_utc = dt_utc.replace(tzinfo=_tz.utc)
            h = (dt_utc + EAT).hour
            for s, e, lbl, bg, rgb in BANDS:
                if s <= h < e: return lbl, bg, rgb
        except Exception: pass
        return "🌙  LATE NIGHT", "051B35", RGB_SKY

    def _eat_time(dt_utc):
        try: return (dt_utc + EAT).strftime("%H:%M")
        except: return ""

    def _eat_date(dt_utc):
        try: return (dt_utc + EAT).strftime("%d %b")
        except: return ""

    def _match_dt(m):
        try:
            dt = _dt.fromisoformat(m.get("start_time", "").replace("Z", "+00:00"))
            if dt.tzinfo is None: dt = dt.replace(tzinfo=_tz.utc)
            return dt
        except Exception: return _now_dt

    # ── XML Helpers ───────────────────────────────────────────────────────────
    def _shd(cell, hex_color):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

    def _margins(cell, top=50, bottom=50, left=70, right=70):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for nm, v in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            n = OxmlElement(f'w:{nm}')
            n.set(qn('w:w'), str(v)); n.set(qn('w:type'), 'dxa')
            tcMar.append(n)
        tcPr.append(tcMar)

    def _cell_borders(cell, color=HEX_BORDER, sz="2"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for nm in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{nm}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
            b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
            tcB.append(b)
        tcPr.append(tcB)

    def _no_borders(table):
        tblPr = table._tbl.tblPr
        tcB = OxmlElement('w:tblBorders')
        for nm in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{nm}'); b.set(qn('w:val'), 'none')
            tcB.append(b)
        tblPr.append(tcB)

    def _ct(cell, text, bold=False, italic=False, color=None, size=None,
            align=None, sa=0, sb=0):
        """Write styled text into a cell, reusing existing empty paragraph."""
        if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after  = Pt(sa)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.name = FF
        if color: r.font.color.rgb = color
        if size:  r.font.size = size
        return p

    def _get_odd(val):
        if val is None: return None
        if isinstance(val, (int, float)): return float(val)
        if isinstance(val, dict):
            for f in ("price", "odd", "odds", "value"):
                if val.get(f):
                    try: return float(val[f])
                    except: pass
        try: return float(val)
        except: return None

    # ── Data Loading ──────────────────────────────────────────────────────────
    sp_available = True
    from app.workers.celery_tasks import _redis as _get_redis_for_sp
    _r_sp = _get_redis_for_sp()

    def _read_sp_raw(sport_slug):
        for key in (f"odds:sp:upcoming:{sport_slug}", f"sp:upcoming:{sport_slug}"):
            try:
                raw = _r_sp.get(key)
                if not raw: continue
                obj = _json.loads(raw)
                if isinstance(obj, list): return [m for m in obj if isinstance(m, dict)]
                if isinstance(obj, dict):
                    ms = obj.get("matches") or obj.get("data") or []
                    if isinstance(ms, list): return [m for m in ms if isinstance(m, dict)]
            except Exception: pass
        return []

    sp_raw_list = _read_sp_raw(sport)
    sp_available = len(sp_raw_list) > 0
    if not sp_raw_list:
        try:
            from app.workers.tasks_upcoming import sp_harvest_sport
            sp_harvest_sport.delay(sport)
        except Exception: pass

    matches_raw = []
    _now_ts = time.time()
    try:
        from app.api.odds_stream import _get_unified_patched
        raw_upcoming = _get_unified_patched("upcoming", sport, force_refresh=False)
        raw_live     = _get_unified_patched("live",     sport, force_refresh=False)
        live_jks = set(); seen_jks = set()
        for m in raw_live:
            if not isinstance(m, dict): continue
            jk = m.get("join_key") or m.get("parent_match_id")
            if jk: live_jks.add(jk)
        for m in raw_live:
            if not isinstance(m, dict): continue
            jk = m.get("join_key") or m.get("parent_match_id")
            if jk and jk not in seen_jks:
                st_raw = m.get("start_time", "")
                if st_raw:
                    try:
                        m_dt2 = _dt.fromisoformat(st_raw.replace("Z", "+00:00"))
                        if m_dt2.tzinfo is None: m_dt2 = m_dt2.replace(tzinfo=_tz.utc)
                        if start_dt and m_dt2 < start_dt: continue
                        if end_dt   and m_dt2 > end_dt:   continue
                    except Exception: pass
                seen_jks.add(jk)
                m.setdefault("_is_live_doc", True)
                if m.get("arb_opportunities") and not m.get("arbitrage"):
                    m["arbitrage"] = m["arb_opportunities"]; m["has_arb"] = True
                matches_raw.append(m)
        for m in raw_upcoming:
            if not isinstance(m, dict): continue
            jk = m.get("join_key") or m.get("parent_match_id")
            if jk in live_jks or (jk and jk in seen_jks): continue
            st_raw = m.get("start_time", "")
            if st_raw:
                try:
                    st_dt2 = _dt.fromisoformat(st_raw.replace("Z", "+00:00"))
                    if st_dt2.tzinfo is None: st_dt2 = st_dt2.replace(tzinfo=_tz.utc)
                    if st_dt2.timestamp() < _now_ts - 90: continue
                    if start_dt and st_dt2 < start_dt: continue
                    if end_dt   and st_dt2 > end_dt:   continue
                except Exception: pass
            if jk: seen_jks.add(jk)
            if m.get("arb_opportunities") and not m.get("arbitrage"):
                m["arbitrage"] = m["arb_opportunities"]; m["has_arb"] = True
            matches_raw.append(m)
    except Exception: pass

    # SP enrichment
    try:
        from app.api.odds_stream import _normalise_markets, _get_price
        _sp_by_br = {}; _sp_by_name = {}
        for sp_m in sp_raw_list:
            br = str(sp_m.get("betradar_id") or "").strip()
            if br: _sp_by_br[br] = sp_m
            h = (sp_m.get("home_team") or "").lower().strip()[:10]
            a = (sp_m.get("away_team") or "").lower().strip()[:10]
            if h and a: _sp_by_name[f"{h}|{a}"] = sp_m

        def _find_sp(um):
            br = str(um.get("betradar_id") or
                     um.get("join_key", "").replace("br_", "") or "").strip()
            if br and br in _sp_by_br: return _sp_by_br[br]
            h = (um.get("home_team") or "").lower().strip()[:10]
            a = (um.get("away_team") or "").lower().strip()[:10]
            if h and a: return _sp_by_name.get(f"{h}|{a}")

        for m in matches_raw:
            sp_m = _find_sp(m)
            if sp_m is None: continue
            if not m.get("sms_id")    and sp_m.get("sms_id"):    m["sms_id"]    = sp_m["sms_id"]
            if not m.get("sp_game_id") and sp_m.get("sp_game_id"): m["sp_game_id"] = sp_m["sp_game_id"]
            sp_mkts = sp_m.get("markets") or {}
            if sp_mkts:
                m.setdefault("bookmakers", {})
                if "sp" not in m["bookmakers"]:
                    m["bookmakers"]["sp"] = {"bookmaker": "SportPesa", "slug": "sp", "markets": {}}
                existing = m["bookmakers"]["sp"].get("markets") or {}
                if len(existing) < len(sp_mkts): m["bookmakers"]["sp"]["markets"] = sp_mkts
            if sp_mkts and m.get("best") is not None:
                try:
                    norm_sp = _normalise_markets(sp_mkts)
                    for mkt, outcomes in norm_sp.items():
                        if not isinstance(outcomes, dict): continue
                        m["best"].setdefault(mkt, {})
                        for out, p in outcomes.items():
                            price = _get_price(p)
                            if price > 1.0:
                                existing = m["best"][mkt].get(out)
                                if not existing or price > existing.get("odd", 0):
                                    m["best"][mkt][out] = {"odd": price, "bk": "sp"}
                except Exception: pass
    except Exception: pass

    if not matches_raw:
        try:
            from app.api.odds_stream import _load_db_matches
            db_up, _, _ = _load_db_matches(sport, mode="upcoming", page=1, per_page=500)
            db_lv, _, _ = _load_db_matches(sport, mode="live",     page=1, per_page=100)
            matches_raw.extend(db_up + db_lv)
        except Exception: pass

    # Arbitrage recalculation
    for m in matches_raw:
        if not m.get("has_arb") or not m.get("arbitrage"):
            best = m.get("best") or {}
            arb_markets_calc = []
            if len(m.get("bookmakers") or {}) >= 2:
                try:
                    from app.workers.arb_engine import detect_arb_for_stream
                    has_arb_calc, _best_arb_calc, arb_markets_calc = detect_arb_for_stream(best)
                    if not has_arb_calc:
                        arb_markets_calc = []
                except Exception:
                    arb_markets_calc = []
            if arb_markets_calc:
                m["has_arb"] = True
                m["arbitrage"] = arb_markets_calc
                m["best_arb_pct"] = arb_markets_calc[0]["profit_pct"]

    # Time-window re-filter
    if start_dt or end_dt:
        filtered = []
        for m in matches_raw:
            st_raw = m.get("start_time", "")
            if not st_raw: continue
            try:
                m_dt3 = _dt.fromisoformat(st_raw.replace("Z", "+00:00"))
                if m_dt3.tzinfo is None: m_dt3 = m_dt3.replace(tzinfo=_tz.utc)
                if start_dt and m_dt3 < start_dt: continue
                if end_dt   and m_dt3 > end_dt:   continue
                filtered.append(m)
            except Exception: pass
        matches_raw = filtered

    matches = matches_raw
    if arb_only: matches = [m for m in matches if m.get("has_arb") or m.get("arbitrage")]
    matches.sort(key=lambda x: x.get("start_time") or "")

    # Build bookmaker game-ID map
    link_dict = {}
    br_ids = [m.get("betradar_id") or m.get("parent_match_id")
              for m in matches if m.get("betradar_id") or m.get("parent_match_id")]
    if br_ids:
        try:
            from app.models.odds import UnifiedMatch
            from app.models.bookmakers_model import BookmakerMatchLink, Bookmaker
            ums = UnifiedMatch.query.filter(UnifiedMatch.parent_match_id.in_(br_ids)).all()
            um_id_to_br = {um.id: um.parent_match_id for um in ums}
            if um_id_to_br:
                bml_list = BookmakerMatchLink.query.filter(
                    BookmakerMatchLink.match_id.in_(list(um_id_to_br.keys()))).all()
                bks    = Bookmaker.query.all()
                bk_map_lnk = {b.id: b.slug for b in bks}
                for bml in bml_list:
                    slug  = bk_map_lnk.get(bml.bookmaker_id, str(bml.bookmaker_id))
                    br_id = um_id_to_br.get(bml.match_id)
                    if br_id: link_dict.setdefault(br_id, {})[slug] = bml.external_match_id
        except Exception: pass

    def _get_ids(m):
        ids = {}
        for slug in ("sp", "bt", "od"):
            val = m.get("sms_id") if slug == "sp" else None
            if not val: val = m.get("sp_game_id") if slug == "sp" else None
            if not val: val = (m.get("bk_ids") or {}).get(slug)
            if not val:
                br_id = m.get("betradar_id") or m.get("parent_match_id")
                if br_id: val = (link_dict.get(br_id) or {}).get(slug)
            if val and str(val).isdigit() and len(str(val)) <= 8:
                ids[slug] = str(val)
        return ids

    # ── Document Setup — Landscape A4, tight margins ──────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.orientation   = WD_ORIENT.LANDSCAPE
        sec.page_width    = Cm(29.7)
        sec.page_height   = Cm(21.0)
        sec.top_margin    = Cm(0.85)
        sec.bottom_margin = Cm(0.75)
        sec.left_margin   = Cm(0.9)
        sec.right_margin  = Cm(0.9)

    doc.styles['Normal'].font.name       = FF
    doc.styles['Normal'].font.size       = PT_TEXT
    doc.styles['Normal'].font.color.rgb  = RGB_TXT

    sport_emojis = {
        "soccer": "⚽ SOCCER", "esoccer": "⚽ ESOCCER",
        "basketball": "🏀 BASKETBALL", "tennis": "🎾 TENNIS",
        "ice-hockey": "🏒 ICE HOCKEY", "volleyball": "🏐 VOLLEYBALL",
        "cricket": "🏏 CRICKET", "rugby": "🏉 RUGBY",
        "table-tennis": "🏓 TABLE TENNIS", "handball": "🤾 HANDBALL",
        "baseball": "⚾ BASEBALL", "mma": "🥊 MMA",
        "boxing": "🥊 BOXING", "darts": "🎯 DARTS",
        "american-football": "🏈 AMERICAN FOOTBALL",
    }
    sport_title = sport_emojis.get(sport.lower(), "🏆 " + sport.upper())
    doc_type    = "ARBITRAGE REPORT" if arb_only else "ODDS BOOKLET"
    USABLE_W    = Cm(27.9)   # 29.7 - 0.9*2

    # ── Title Bar ─────────────────────────────────────────────────────────────
    tb = doc.add_table(rows=1, cols=2)
    _no_borders(tb)
    tb.autofit = False
    tb.columns[0].width = Cm(19.0)
    tb.columns[1].width = Cm(8.9)
    tl, tr = tb.rows[0].cells[0], tb.rows[0].cells[1]
    _shd(tl, HEX_PRI); _shd(tr, HEX_PRI)
    _margins(tl, top=130, bottom=130, left=200, right=100)
    _margins(tr, top=130, bottom=130, left=100, right=200)

    # Left: sport title + type
    p_t = tl.paragraphs[0]; p_t.paragraph_format.space_after = Pt(3); p_t.paragraph_format.line_spacing = 1.0
    rt1 = p_t.add_run(f"{sport_title}  "); rt1.bold = True; rt1.font.size = PT_TITLE; rt1.font.color.rgb = RGB_W; rt1.font.name = FF
    rt2 = p_t.add_run(doc_type);           rt2.bold = True; rt2.font.size = PT_TITLE; rt2.font.color.rgb = RGB_SKY; rt2.font.name = FF
    p_s = tl.add_paragraph(); p_s.paragraph_format.space_after = Pt(0); p_s.paragraph_format.line_spacing = 1.0
    eat_now = _now_dt + EAT
    rs = p_s.add_run(f"🗓  {eat_now.strftime('%A, %d %B %Y  %H:%M EAT')}   |   📋 {len(matches)} {'match' if len(matches)==1 else 'matches'}")
    rs.font.size = PT_SUB; rs.font.color.rgb = RGB_MUT; rs.font.name = FF

    # Right: bookmaker legend
    p_l = tr.paragraphs[0]; p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_l.paragraph_format.space_after = Pt(4); p_l.paragraph_format.line_spacing = 1.0
    for badge, bclr, lbl, sep in [("● SP", RGB_SP, " SportPesa", "   "),
                                    ("● BT", RGB_BT, " Betika", "   "),
                                    ("● OD", RGB_OD, " OdiBets", "")]:
        rb = p_l.add_run(badge); rb.bold = True; rb.font.color.rgb = bclr; rb.font.size = PT_TINY; rb.font.name = FF
        rl = p_l.add_run(lbl + sep); rl.font.size = PT_TINY; rl.font.color.rgb = RGB_W; rl.font.name = FF
    p_l2 = tr.add_paragraph(); p_l2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_l2.paragraph_format.space_after = Pt(0); p_l2.paragraph_format.line_spacing = 1.0
    rl2 = p_l2.add_run("🟢 Green cell = best available odd   |   IDs = bookmaker game codes")
    rl2.font.size = Pt(5.5); rl2.font.color.rgb = RGB_MUT; rl2.font.name = FF

    # Accent stripe under title
    p_acc = doc.add_paragraph(); p_acc.paragraph_format.space_after = Pt(6); p_acc.paragraph_format.space_before = Pt(0)
    r_acc = p_acc.add_run("▬" * 260); r_acc.font.size = Pt(1.5); r_acc.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8); r_acc.font.name = FF

    if not matches:
        pn = doc.add_paragraph(); pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pn.add_run("No matches found for the selected filters.").italic = True
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf, sp_available

    # ═══════════════════════════════════════════════════════════════════════
    # ── ARBITRAGE-ONLY DOCUMENT ──────────────────────────────────────────
    # ═══════════════════════════════════════════════════════════════════════
    if arb_only:
        prev_band = None
        for global_idx, m in enumerate(matches, 1):
            arbs = m.get("arbitrage") or []
            if not arbs: continue
            m_dt = _match_dt(m)
            band_lbl, band_hex, band_rgb = _band_info(m_dt)

            # Time-band sub-header
            if band_lbl != prev_band:
                prev_band = band_lbl
                bt = doc.add_table(rows=1, cols=1); bt.autofit = False; _no_borders(bt)
                bt.columns[0].width = USABLE_W
                bc = bt.rows[0].cells[0]; _shd(bc, band_hex); _margins(bc, top=70, bottom=70, left=150, right=150)
                p_bd = bc.paragraphs[0]; p_bd.paragraph_format.line_spacing = 1.0
                rb = p_bd.add_run(f"  {band_lbl}  "); rb.bold = True; rb.font.size = PT_BAND; rb.font.color.rgb = band_rgb; rb.font.name = FF
                sp_after = doc.add_paragraph(); sp_after.paragraph_format.space_after = Pt(4); sp_after.paragraph_format.space_before = Pt(8)

            h_team = (m.get("home_team") or "Home")[:28]
            a_team = (m.get("away_team") or "Away")[:28]
            comp   = m.get("competition") or ""
            ids    = _get_ids(m)

            # Match header (deep green card)
            mht = doc.add_table(rows=1, cols=1); mht.autofit = False; _no_borders(mht)
            mht.columns[0].width = USABLE_W
            mhc = mht.rows[0].cells[0]; _shd(mhc, "064E3B"); _margins(mhc, top=100, bottom=80, left=160, right=160)
            p_mh = mhc.paragraphs[0]; p_mh.paragraph_format.line_spacing = 1.0; p_mh.paragraph_format.space_after = Pt(2)
            r_num  = p_mh.add_run(f"{global_idx}.  "); r_num.font.size = PT_TEXT; r_num.font.color.rgb = RGB_ARB; r_num.font.name = FF
            r_mname= p_mh.add_run(f"{h_team}  vs  {a_team}"); r_mname.bold = True; r_mname.font.size = Pt(9.5); r_mname.font.color.rgb = RGB_W; r_mname.font.name = FF
            r_ko   = p_mh.add_run(f"    ⏱ {_eat_date(m_dt)} {_eat_time(m_dt)}"); r_ko.font.size = PT_TINY; r_ko.font.color.rgb = RGB_ARB; r_ko.font.name = FF
            if comp:
                r_comp = p_mh.add_run(f"    🏆 {comp}"); r_comp.font.size = PT_TINY; r_comp.font.color.rgb = RGB_ARB; r_comp.font.name = FF
            if ids:
                p_ids = mhc.add_paragraph(); p_ids.paragraph_format.line_spacing = 1.0; p_ids.paragraph_format.space_after = Pt(0)
                r_ids = p_ids.add_run("📲  " + "   ".join([f"{s.upper()} #{v}" for s, v in ids.items()]))
                r_ids.font.size = PT_TINY; r_ids.font.color.rgb = RGBColor(0xA7, 0xF3, 0xD0); r_ids.font.name = FF

            # Arb opportunities table
            arb_tbl = doc.add_table(rows=1, cols=5); arb_tbl.autofit = False; _no_borders(arb_tbl)
            arb_col_ws = [Cm(4.2), Cm(2.6), Cm(7.0), Cm(7.0), Cm(7.1)]
            for ci, w in enumerate(arb_col_ws): arb_tbl.columns[ci].width = w
            hdr_row = arb_tbl.rows[0]
            for ci, (lbl, aln) in enumerate([
                ("Market",               WD_ALIGN_PARAGRAPH.LEFT),
                ("Profit %",             WD_ALIGN_PARAGRAPH.CENTER),
                ("Leg 1 — Bk / Outcome / Odd",  WD_ALIGN_PARAGRAPH.CENTER),
                ("Leg 2 — Bk / Outcome / Odd",  WD_ALIGN_PARAGRAPH.CENTER),
                ("Stake Distribution (KES 2,000 total)",   WD_ALIGN_PARAGRAPH.LEFT),
            ]):
                c = hdr_row.cells[ci]; _shd(c, "064E3B"); _margins(c, top=80, bottom=80, left=80, right=80)
                _ct(c, lbl, bold=True, color=RGB_ARB, size=PT_HDR, align=aln)

            for ai, arb in enumerate(arbs[:6]):
                mkt_lbl = arb.get("market", "").replace("_", " ").title()
                profit  = float(arb.get("profit_pct", 0))
                legs    = arb.get("legs", [])
                s_inv   = sum(1.0 / float(l.get("odd", 1)) for l in legs if l.get("odd"))
                payout  = (2000.0 / s_inv) if s_inv > 0 else 0.0
                bg      = "F0FDF4" if ai % 2 == 0 else "FFFFFF"
                dr      = arb_tbl.add_row()
                for ci2 in range(5):
                    c2 = dr.cells[ci2]; _shd(c2, bg); _margins(c2, top=60, bottom=60, left=80, right=80); _cell_borders(c2, "BBF7D0", "2")

                _ct(dr.cells[0], mkt_lbl, bold=True, color=RGB_GRN, size=PT_TEXT)
                _ct(dr.cells[1], f"✅ {profit:.2f}%", bold=True, color=RGB_GRN, size=PT_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)

                for li, leg in enumerate(legs[:2]):
                    odd    = float(leg.get("odd", 0) or 0)
                    bk     = str(leg.get("bk", "")).upper()
                    outcome= str(leg.get("outcome", "")).upper()
                    bk_clr = RGB_SP if bk == "SP" else (RGB_BT if bk == "BT" else RGB_OD)
                    c_leg  = dr.cells[2 + li]
                    p_leg  = c_leg.paragraphs[0]; p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_leg.paragraph_format.line_spacing = 1.0
                    rb2 = p_leg.add_run(f"[{bk}]  "); rb2.bold = True; rb2.font.size = PT_TINY; rb2.font.color.rgb = bk_clr; rb2.font.name = FF
                    ro  = p_leg.add_run(f"{outcome}  "); ro.font.size = PT_TINY; ro.font.color.rgb = RGB_TXT; ro.font.name = FF
                    rv  = p_leg.add_run(f"@ {odd:.2f}"); rv.bold = True; rv.font.size = PT_TEXT; rv.font.color.rgb = RGB_GRN; rv.font.name = FF

                # Stake column
                c_st = dr.cells[4]
                stake_parts = []
                for leg in legs:
                    odd = float(leg.get("odd", 1) or 1)
                    bk  = str(leg.get("bk", "")).upper()
                    outcome = str(leg.get("outcome", "")).upper()
                    amt = (1.0 / odd) / s_inv * 2000.0 if s_inv > 0 else 0
                    stake_parts.append(f"[{bk}] {outcome}: KES {amt:.0f}")
                p_st = c_st.paragraphs[0]; p_st.paragraph_format.line_spacing = 1.0
                rst = p_st.add_run("   |   ".join(stake_parts)); rst.font.size = PT_TINY; rst.font.color.rgb = RGB_TXT; rst.font.name = FF
                p_po = c_st.add_paragraph(); p_po.paragraph_format.space_before = Pt(3); p_po.paragraph_format.line_spacing = 1.0
                rpo = p_po.add_run(f"💰  Guaranteed Payout: KES {payout:.0f}   (Profit: {profit:.2f}%)")
                rpo.bold = True; rpo.font.size = PT_TINY; rpo.font.color.rgb = RGB_GRN; rpo.font.name = FF

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ═══════════════════════════════════════════════════════════════════════
    # ── REGULAR BOOKLET — Market-First Grouped Layout ────────────────────
    # ═══════════════════════════════════════════════════════════════════════
    else:
        # Outcome alias map for canonical names → possible dict keys
        _ALIASES = {
            "1":    ["1", "home", "home_win", "win", "home win"],
            "X":    ["x", "draw", "tie"],
            "2":    ["2", "away", "away_win", "loss", "away win"],
            "Yes":  ["yes", "btts_yes", "both_score", "yesbtts"],
            "No":   ["no",  "btts_no",  "nonebtts"],
            "1X":   ["1x", "home_or_draw", "homeordraw"],
            "12":   ["12", "home_or_away", "homeoraway"],
            "X2":   ["x2", "draw_or_away", "draworaway"],
            "Over": ["over", "o"],
            "Under":["under", "u"],
            "Home": ["home", "home_win", "1"],
            "Away": ["away", "away_win", "2"],
        }

        def _best_for_out(m, mkt_keys, out_canon):
            """Return (best_odd_float, bk_slug) for a canonical outcome across market keys."""
            aliases = [a.lower() for a in _ALIASES.get(out_canon, [out_canon.lower()])]
            best_odd = None; best_bk = ""
            for mkt_key in mkt_keys:
                mkt_data = (m.get("best") or {}).get(mkt_key) or {}
                for k, v in mkt_data.items():
                    if str(k).lower() in aliases:
                        fv  = _get_odd(v.get("odd") if isinstance(v, dict) else v) if isinstance(v, (dict, int, float, str)) else None
                        bk  = (v.get("bk", "") if isinstance(v, dict) else "")
                        if fv and (best_odd is None or fv > best_odd):
                            best_odd = fv; best_bk = bk
            return best_odd, best_bk

        def _section_header(title_txt):
            sp_b = doc.add_paragraph(); sp_b.paragraph_format.space_after = Pt(0); sp_b.paragraph_format.space_before = Pt(10)
            ht = doc.add_table(rows=1, cols=1); ht.autofit = False; _no_borders(ht)
            ht.columns[0].width = USABLE_W
            hc = ht.rows[0].cells[0]; _shd(hc, HEX_PRI); _margins(hc, top=120, bottom=120, left=180, right=180)
            p_h = hc.paragraphs[0]; p_h.paragraph_format.line_spacing = 1.0
            rh = p_h.add_run(title_txt); rh.bold = True; rh.font.size = PT_SEC; rh.font.color.rgb = RGB_W; rh.font.name = FF
            # Sky-blue accent line
            pa = doc.add_paragraph(); pa.paragraph_format.space_after = Pt(0); pa.paragraph_format.space_before = Pt(0)
            ra = pa.add_run("▬" * 260); ra.font.size = Pt(1.5); ra.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8); ra.font.name = FF

        def _band_header(band_lbl, band_hex, band_rgb, n_matches):
            bt = doc.add_table(rows=1, cols=2); bt.autofit = False; _no_borders(bt)
            bt.columns[0].width = Cm(20.0); bt.columns[1].width = Cm(7.9)
            bl, br2 = bt.rows[0].cells[0], bt.rows[0].cells[1]
            _shd(bl, band_hex); _shd(br2, band_hex)
            _margins(bl, top=65, bottom=65, left=140, right=100)
            _margins(br2, top=65, bottom=65, left=100, right=140)
            p_bl = bl.paragraphs[0]; p_bl.paragraph_format.line_spacing = 1.0
            rb3 = p_bl.add_run(f"  {band_lbl}  "); rb3.bold = True; rb3.font.size = PT_BAND; rb3.font.color.rgb = band_rgb; rb3.font.name = FF
            p_br3 = br2.paragraphs[0]; p_br3.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_br3.paragraph_format.line_spacing = 1.0
            rb4 = p_br3.add_run(f"{n_matches} match{'es' if n_matches!=1 else ''}  "); rb4.font.size = PT_TINY; rb4.font.color.rgb = band_rgb; rb4.font.name = FF

        def _col_widths(n_out, has_date):
            # Fixed allocation in cm, total must ≤ 27.9
            num_w   = 0.55
            match_w = 5.2
            date_w  = 1.7 if has_date else 0.0
            ko_w    = 1.35
            out_w   = 1.85 if n_out == 2 else 1.65  # per outcome
            bk_w    = 1.7
            ids_used = num_w + match_w + date_w + ko_w + out_w * n_out + bk_w
            ids_w    = max(27.9 - ids_used, 2.0)
            ws = [num_w, match_w]
            if has_date: ws.append(date_w)
            ws += [ko_w] + [out_w] * n_out + [bk_w, ids_w]
            return ws

        def _market_table_header(outcomes, has_date, col_ws):
            n = len(col_ws)
            t = doc.add_table(rows=1, cols=n); t.autofit = False; _no_borders(t)
            for ci, w in enumerate(col_ws): t.columns[ci].width = Cm(w)
            hdr = t.rows[0]
            labels  = ["#", "Home vs Away"]
            if has_date: labels.append("Date")
            labels += ["KO"] + outcomes + ["Best BK", "Game IDs"]
            aligns  = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            if has_date: aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
            aligns += [WD_ALIGN_PARAGRAPH.CENTER] * (len(outcomes) + 1) + [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            for ci, (lbl, aln) in enumerate(zip(labels, aligns)):
                c = hdr.cells[ci]; _shd(c, "1E293B"); _margins(c, top=75, bottom=75, left=55, right=55)
                _ct(c, lbl, bold=True, color=RGB_SKY, size=PT_HDR, align=aln)
            return t

        # Determine if multi-day (needs a Date column)
        has_date_col = False
        if len(matches) >= 2:
            try:
                d0 = _match_dt(matches[0]); d1 = _match_dt(matches[-1])
                has_date_col = (d1 - d0).days >= 1
            except Exception: pass

        # Build index for global numbering
        match_global_idx = {id(m): i + 1 for i, m in enumerate(matches)}

        # Collect distinct Over/Under lines (max 5)
        ou_lines = []
        seen_ou  = set()
        for m in matches:
            for k in (m.get("best") or {}):
                if k.startswith(("over_under_goals_", "over_under_")):
                    raw = k.replace("over_under_goals_", "").replace("over_under_", "").replace("_", ".")
                    try:
                        fv = float(raw)
                        if fv not in seen_ou: seen_ou.add(fv); ou_lines.append(fv)
                    except Exception: pass
        ou_lines = sorted(ou_lines)[:5]

        # Collect distinct Asian Handicap lines (max 3)
        ah_lines = []
        seen_ah  = set()
        for m in matches:
            for k in (m.get("best") or {}):
                if k.startswith("asian_handicap_"):
                    raw = k.replace("asian_handicap_", "").replace("_", ".")
                    if raw not in seen_ah: seen_ah.add(raw); ah_lines.append(raw)
        ah_lines = ah_lines[:3]

        # Master list of sections
        ALL_SECTIONS = [
            ("🏆  FULL-TIME 1X2",      ["1x2", "match_winner", "moneyline"],     ["1", "X", "2"]),
            ("⏱  HALF-TIME RESULT",    ["half_time"],                             ["1", "X", "2"]),
            ("⚽  BOTH TEAMS TO SCORE", ["btts"],                                  ["Yes", "No"]),
            ("🔄  DOUBLE CHANCE",       ["double_chance"],                         ["1X", "12", "X2"]),
            ("🎯  DRAW NO BET",         ["dnb"],                                   ["1", "2"]),
        ]
        for line in ou_lines:
            lk = str(line).replace(".", "_")
            ALL_SECTIONS.append((
                f"📊  OVER / UNDER  {line}  GOALS",
                [f"over_under_goals_{lk}", f"over_under_{lk}"],
                ["Over", "Under"]
            ))
        for ah in ah_lines:
            ALL_SECTIONS.append((
                f"📐  ASIAN HANDICAP  {ah}",
                [f"asian_handicap_{ah.replace('.','_')}"],
                ["Home", "Away"]
            ))

        for sec_title, mkt_keys, outcomes in ALL_SECTIONS:
            # Filter matches with data in this market
            sec_matches = [m for m in matches
                           if any((m.get("best") or {}).get(k) for k in mkt_keys)]
            if not sec_matches: continue

            _section_header(sec_title)
            col_ws  = _col_widths(len(outcomes), has_date_col)
            prev_bl  = None
            curr_tbl = None
            row_idx  = 0

            for m in sec_matches:
                m_dt = _match_dt(m)
                band_lbl, band_hex, band_rgb = _band_info(m_dt)

                # New time band → add sub-header + fresh table header
                if band_lbl != prev_bl:
                    prev_bl  = band_lbl
                    # Count how many sec_matches fall in this band
                    n_band = sum(1 for x in sec_matches if _band_info(_match_dt(x))[0] == band_lbl)
                    _band_header(band_lbl, band_hex, band_rgb, n_band)
                    curr_tbl = _market_table_header(outcomes, has_date_col, col_ws)
                    row_idx  = 0

                h_team = (m.get("home_team") or "Home")[:22]
                a_team = (m.get("away_team") or "Away")[:22]
                ko_str = _eat_time(m_dt)
                dt_str = _eat_date(m_dt)
                ids    = _get_ids(m)
                ids_str = "  ".join([f"{s.upper()}#{v}" for s, v in ids.items()])

                # Resolve best odd per outcome
                best_outs = [_best_for_out(m, mkt_keys, out) for out in outcomes]
                # Find max odd across all outcomes for highlighting
                max_odd = max((o for o, _ in best_outs if o), default=None)
                # Unique BKs that hold the best odds (for the "Best BK" column)
                bk_set = []; bk_seen = set()
                for _, bk in best_outs:
                    if bk and bk.upper() not in bk_seen:
                        bk_seen.add(bk.upper()); bk_set.append(bk.upper())
                bk_summary = " / ".join(bk_set) or "—"

                # Add data row
                dr = curr_tbl.add_row()
                bg_hex = HEX_ALT if row_idx % 2 == 1 else "FFFFFF"
                for ci2 in range(len(col_ws)):
                    c2 = dr.cells[ci2]; _shd(c2, bg_hex)
                    _margins(c2, top=40, bottom=40, left=55, right=55)
                    _cell_borders(c2, HEX_BORDER, "2")

                col_i = 0
                global_n = match_global_idx.get(id(m), 0)
                _ct(dr.cells[col_i], str(global_n), color=RGB_MUT, size=PT_TINY, align=WD_ALIGN_PARAGRAPH.CENTER); col_i += 1

                # Match name (home bold, "v" muted, away normal)
                mc = dr.cells[col_i]; col_i += 1
                pm = mc.paragraphs[0]; pm.paragraph_format.line_spacing = 1.0
                rh2 = pm.add_run(h_team); rh2.bold = True; rh2.font.size = PT_TEXT; rh2.font.color.rgb = RGB_TXT; rh2.font.name = FF
                rv2 = pm.add_run("  v  "); rv2.font.size = PT_TINY; rv2.font.color.rgb = RGB_MUT; rv2.font.name = FF
                ra2 = pm.add_run(a_team); ra2.font.size = PT_TEXT; ra2.font.color.rgb = RGB_TXT; ra2.font.name = FF

                if has_date_col:
                    _ct(dr.cells[col_i], dt_str, color=RGB_MUT, size=PT_TINY, align=WD_ALIGN_PARAGRAPH.CENTER); col_i += 1

                _ct(dr.cells[col_i], ko_str, bold=True, color=RGB_TXT, size=PT_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER); col_i += 1

                for b_odd, b_bk in best_outs:
                    oc = dr.cells[col_i]; col_i += 1
                    if b_odd and b_odd > 1.0:
                        is_max = (b_odd == max_odd)
                        if is_max: _shd(oc, HEX_BEST)
                        po = oc.paragraphs[0]; po.alignment = WD_ALIGN_PARAGRAPH.CENTER; po.paragraph_format.line_spacing = 1.0
                        rv3 = po.add_run(f"{b_odd:.2f}"); rv3.bold = is_max; rv3.font.size = PT_TEXT
                        rv3.font.color.rgb = RGB_GRN if is_max else RGB_TXT; rv3.font.name = FF
                        if b_bk:
                            bk_clr2 = RGB_SP if b_bk.lower() == "sp" else (RGB_BT if b_bk.lower() == "bt" else RGB_OD)
                            rb5 = po.add_run(f"\n{b_bk.upper()}"); rb5.font.size = Pt(5.5); rb5.font.color.rgb = bk_clr2; rb5.font.name = FF
                    else:
                        _ct(oc, "—", color=RGB_MUT, size=PT_TINY, align=WD_ALIGN_PARAGRAPH.CENTER)

                bk_clr_main = RGB_SP if "SP" in bk_set else (RGB_BT if "BT" in bk_set else RGB_OD)
                _ct(dr.cells[col_i], bk_summary, bold=True, color=bk_clr_main, size=PT_TINY, align=WD_ALIGN_PARAGRAPH.CENTER); col_i += 1
                _ct(dr.cells[col_i], ids_str, color=RGB_MUT, size=PT_TINY); col_i += 1

                row_idx += 1

            # Spacer after each market section
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # ── Footer ────────────────────────────────────────────────────────────────
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.paragraph_format.space_before = Pt(14)
    rf = pf.add_run(
        f"📊 OddsKenya  |  All odds sourced live from SP / BT / OD — verify before placing bets.  |  "
        f"Generated {eat_now.strftime('%d %b %Y %H:%M')} EAT"
    )
    rf.font.size = PT_TINY; rf.italic = True; rf.font.color.rgb = RGB_MUT; rf.font.name = FF

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, sp_available


def download_odds_word():
    from app.utils.customer_jwt_helpers import _current_user_from_header
    from flask import send_file, make_response
    import time

    sport    = request.args.get("sport", "soccer").lower().strip()
    arb_only = request.args.get("arb_only", "") in ("1", "true")
    preset   = request.args.get("preset", "all").lower().strip()

    # 1. AUTH RULE: Only soccer (football) is free & anonymous.
    #    Other sports require an authenticated session.
    user = None
    if sport != "soccer":
        user = _current_user_from_header()
        if not user:
            return make_response("Authentication required to download reports for this sport.", 401)

    # Log report download for monetization funnel analytics
    from app.utils.decorators_ import log_event
    log_event("report_download", {"sport": sport, "arb_only": arb_only, "preset": preset})

    sp_available = True
    # 2. Try serving the pre-generated MinIO-cached document first (fast path)
    f_stream = _serve_minio_report(sport, arb_only, preset=preset)
    if f_stream is not None:
        from app.workers.celery_tasks import _redis as _get_redis
        r = _get_redis()
        sp_flag = r.get(f"odds_report:sp_available:{sport}:{preset}")
        if sp_flag is not None:
            sp_available = (sp_flag == b"1" or sp_flag == "1")

    # 3. Fall back to on-demand generation if MinIO is unavailable or cache is stale
    if f_stream is None:
        from datetime import datetime as _dt, timezone as _tz, timedelta
        now = _dt.now(_tz.utc)
        start_dt = None
        end_dt = None
        if preset == "today":
            start_dt = now
            eat_now = now + timedelta(hours=3)
            eat_end = eat_now.replace(hour=23, minute=59, second=59, microsecond=999999)
            end_dt = eat_end - timedelta(hours=3)
        elif preset == "tomorrow":
            eat_now = now + timedelta(hours=3)
            eat_tomorrow_start = (eat_now + timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)
            eat_tomorrow_end = (eat_now + timedelta(days=1)).replace(hour=23, minute=59, second=59, microsecond=999999)
            start_dt = eat_tomorrow_start - timedelta(hours=3)
            end_dt = eat_tomorrow_end - timedelta(hours=3)
        elif preset == "week":
            start_dt = now
            end_dt = now + timedelta(days=7)
        elif preset == "month":
            start_dt = now
            end_dt = now + timedelta(days=30)

        start_time_str = start_dt.strftime("%Y-%m-%dT%H:%M:%SZ") if start_dt else None
        end_time_str = end_dt.strftime("%Y-%m-%dT%H:%M:%SZ") if end_dt else None

        f_stream, sp_available = _generate_word_document(sport, arb_only, start_time_str=start_time_str, end_time_str=end_time_str)
        # Persist to MinIO in the background so the next request is instant
        try:
            _save_minio_report(sport, arb_only, f_stream, preset=preset)
            from app.workers.celery_tasks import _redis as _get_redis
            r = _get_redis()
            r.set(f"odds_report:sp_available:{sport}:{preset}", "1" if sp_available else "0", ex=86400)
        except Exception:
            pass

    filename = f"OddsKenya_Report_{sport}_{time.strftime('%Y%m%d_%H%M%S')}.docx"
    response = make_response(send_file(
        f_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    ))
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["X-SportPesa-Available"] = "1" if sp_available else "0"
    return response


@bp_odds_customer.route("/odds/download/word/async", methods=["POST"])
def download_odds_word_async():
    from app.utils.customer_jwt_helpers import _current_user_from_header
    from flask import jsonify, request
    import logging

    # Parse JSON body
    data = request.get_json(silent=True) or {}
    sport = data.get("sport", "soccer").lower().strip()
    arb_only = bool(data.get("arb_only", False))
    start_time = data.get("start_time")
    end_time = data.get("end_time")
    preset = data.get("preset", "all").lower().strip()

    # AUTH RULE: Only soccer is free. Other sports require authorization.
    if sport != "soccer":
        user = _current_user_from_header()
        if not user:
            return jsonify({"error": "Authentication required to download reports for this sport."}), 401

    # Import the Celery task
    try:
        from app.workers.tasks_ops import generate_custom_report
        task = generate_custom_report.delay(sport, arb_only, start_time, end_time, preset=preset)
        return jsonify({"task_id": task.id, "status": "PENDING"}), 202
    except Exception as exc:
        logging.error("Failed to enqueue custom report task: %s", exc)
        return jsonify({"error": f"Failed to enqueue background export task: {str(exc)}"}), 500


@bp_odds_customer.route("/odds/download/word/status/<task_id>", methods=["GET"])
def download_odds_word_status(task_id):
    from app.workers.celery_tasks import celery
    from flask import jsonify
    
    # Query task status from Celery
    res = celery.AsyncResult(task_id)
    state = res.state  # PENDING, STARTED, RETRY, SUCCESS, FAILURE
    
    # Also check if it's already generated and stored in Redis
    from app.workers.celery_tasks import _redis as _get_redis
    r = _get_redis()
    redis_key = f"custom_report:{task_id}"
    exists = r.exists(redis_key)
    
    if exists:
        return jsonify({"status": "SUCCESS", "task_id": task_id})
    
    if state == "SUCCESS" and not exists:
        # Task says success, but data not in redis (maybe expired?)
        return jsonify({"status": "EXPIRED", "task_id": task_id})
    elif state in ("FAILURE", "REVOKED"):
        return jsonify({"status": "FAILED", "task_id": task_id, "error": str(res.result or "Task failed")})
    
    return jsonify({"status": state, "task_id": task_id})


@bp_odds_customer.route("/odds/download/word/retrieve/<task_id>", methods=["GET"])
def download_odds_word_retrieve(task_id):
    from flask import send_file, make_response, jsonify
    from app.workers.celery_tasks import _redis as _get_redis
    import base64
    import io
    import time
    
    r = _get_redis()
    redis_key = f"custom_report:{task_id}"
    encoded = r.get(redis_key)
    if not encoded:
        return jsonify({"error": "Report not found or expired. Please generate a new report."}), 404
        
    try:
        data_bytes = base64.b64decode(encoded)
        f_stream = io.BytesIO(data_bytes)
        f_stream.seek(0)
        
        filename = f"OddsKenya_Custom_Report_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        response = make_response(send_file(
            f_stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        ))
        response.headers["Access-Control-Allow-Origin"] = "*"
        return response
    except Exception as exc:
        return jsonify({"error": f"Failed to retrieve report: {str(exc)}"}), 500