"""
app/views/customer/word_generator_v2.py
────────────────────────────────────────
Per-time-group Word document generator.

DATA SOURCE — DB first, Redis fallback:
  Primary:  PostgreSQL (UnifiedMatch + BookmakerMatchOdds)
            → works for ALL sports regardless of Redis state
  Fallback: Redis unified cache
            → used only when DB has no matches for the requested window

This fixes "soccer only" availability: soccer was always in Redis because
SP harvests it frequently.  Non-soccer sports are harvested less often so
Redis keys may expire.  The DB is always populated by _persist_bk_matches /
_upsert_and_chain after every harvest.

Public API
──────────
    get_available_groups(sport, date_str)  → list[GroupInfo]
    generate_group_document(sport, group_id, date_str, market_filter) → BytesIO
"""

from __future__ import annotations

import io
from datetime import datetime, timezone, timedelta
from typing import Optional

# ── EAT offset ────────────────────────────────────────────────────────────────
EAT = timedelta(hours=3)

# ── Time groups (EAT hours) ────────────────────────────────────────────────────
TIME_GROUPS: dict[str, tuple[str, int, int]] = {
    "late_night_early": ("🌙 Late Night / Early Hours",  0,  6),
    "early_morning":    ("🌅 Early Morning Kickoffs",    6,  10),
    "morning":          ("☀️  Morning Kickoffs",          10, 14),
    "afternoon":        ("🌤  Afternoon Kickoffs",        14, 18),
    "evening":          ("🌆  Evening Kickoffs",          18, 21),
    "late_night":       ("🌙  Late Night Kickoffs",       21, 24),
}

# ── Bookmaker config ───────────────────────────────────────────────────────────
PRIMARY_BKS = ["sp", "bt", "od"]

BK_LABELS: dict[str, str] = {
    "sp": "SP", "bt": "BT", "od": "OD",
    "1xbet": "1X", "22bet": "22B", "betwinner": "BW",
    "melbet": "MB", "megapari": "MP", "helabet": "HB",
    "paripesa": "PP", "sbo": "SBO",
}

BK_COLORS: dict[str, tuple[str, str]] = {
    "sp":        ("2563EB", "DBEAFE"),
    "bt":        ("16A34A", "DCFCE7"),
    "od":        ("D97706", "FEF3C7"),
    "1xbet":     ("7C3AED", "EDE9FE"),
    "22bet":     ("0891B2", "CFFAFE"),
    "betwinner": ("BE185D", "FCE7F3"),
    "melbet":    ("B45309", "FEF3C7"),
}

# ── Sport slug → DB name variants ─────────────────────────────────────────────
_SPORT_DB_NAMES: dict[str, list[str]] = {
    "soccer":            ["Soccer", "Football"],
    "basketball":        ["Basketball"],
    "tennis":            ["Tennis"],
    "ice-hockey":        ["Ice Hockey"],
    "volleyball":        ["Volleyball"],
    "cricket":           ["Cricket"],
    "rugby":             ["Rugby"],
    "table-tennis":      ["Table Tennis"],
    "handball":          ["Handball"],
    "mma":               ["MMA"],
    "boxing":            ["Boxing"],
    "darts":             ["Darts"],
    "american-football": ["American Football"],
    "baseball":          ["Baseball"],
    "esoccer":           ["eSoccer", "eFootball", "ESoccer"],
}

# ── Bookmaker name → slug ──────────────────────────────────────────────────────
_BK_NAME_TO_SLUG: dict[str, str] = {
    "sportpesa": "sp", "sport pesa": "sp", "sp": "sp",
    "betika": "bt", "bt": "bt",
    "odibets": "od", "odi": "od", "od": "od",
    "1xbet": "1xbet", "22bet": "22bet",
    "betwinner": "betwinner", "melbet": "melbet",
    "megapari": "megapari", "helabet": "helabet",
    "paripesa": "paripesa", "sbo": "sbo",
}

# ─────────────────────────────────────────────────────────────────────────────
# SPORT-SPECIFIC MARKET DEFINITIONS
# Each entry: (display_label, [market_key_aliases], [(out_label, [out_aliases])])
# ─────────────────────────────────────────────────────────────────────────────

# Shared outcome aliases reused across sports
_WIN_DRAW_WIN = [
    ("1", ["1", "home", "home_win", "win", "w1"]),
    ("X", ["x", "draw", "tie"]),
    ("2", ["2", "away", "away_win", "loss", "w2"]),
]
_WIN_LOSE = [
    ("1", ["1", "home", "home_win", "win", "w1"]),
    ("2", ["2", "away", "away_win", "loss", "w2"]),
]
_OVER_UNDER = [
    ("Over",  ["over", "o", "ov"]),
    ("Under", ["under", "u", "un"]),
]
_YES_NO = [
    ("Yes", ["yes", "y", "gg"]),
    ("No",  ["no",  "n", "ng"]),
]

_SPORT_MARKET_DEFS: dict[str, list] = {

    # ── Soccer ────────────────────────────────────────────────────────────────
    "soccer": [
        ("Full-Time 1X2",
         ["1x2", "match_winner", "moneyline", "full_time_result"],
         _WIN_DRAW_WIN),
        ("Both Teams to Score",
         ["btts", "both_teams_to_score", "gg_ng"],
         _YES_NO),
        ("Double Chance",
         ["double_chance"],
         [("1X", ["1x", "home_or_draw"]),
          ("12", ["12", "home_or_away"]),
          ("X2", ["x2", "draw_or_away"])]),
        ("Draw No Bet",
         ["dnb", "draw_no_bet"],
         _WIN_LOSE),
        ("Half-Time 1X2",
         ["half_time", "half_time_result", "ht_result"],
         _WIN_DRAW_WIN),
        ("Over 2.5 Goals",
         ["over_under_goals_2_5", "over_under_2_5", "ou_2_5"],
         _OVER_UNDER),
        ("Over 1.5 Goals",
         ["over_under_goals_1_5", "over_under_1_5", "ou_1_5"],
         _OVER_UNDER),
        ("Over 3.5 Goals",
         ["over_under_goals_3_5", "over_under_3_5", "ou_3_5"],
         _OVER_UNDER),
    ],

    # ── Basketball ────────────────────────────────────────────────────────────
    "basketball": [
        ("Match Winner",
         ["match_winner", "moneyline", "basketball_1x2", "1x2",
          "basketball_match_winner", "outright_winner"],
         _WIN_LOSE),
        ("Match Winner incl. OT",
         ["match_winner_incl_ot", "win_including_ot", "winner_incl_ot"],
         _WIN_LOSE),
        ("Half-Time Winner",
         ["first_half_1x2", "half_time", "first_half_winner",
          "basketball_first_half", "h1_winner"],
         _WIN_LOSE),
        ("Quarter Winner (1st)",
         ["quarter_winner", "first_quarter_winner", "q1_winner"],
         _WIN_LOSE),
        ("Over/Under Points",
         ["over_under", "total_points", "total", "over_under_points",
          "basketball_over_under"],
         _OVER_UNDER),
        ("Asian Handicap",
         ["asian_handicap", "point_spread", "handicap",
          "basketball_asian_handicap"],
         [("Home", ["home", "1", "h"]),
          ("Away", ["away", "2", "a"])]),
        ("1st Half Over/Under",
         ["first_half_over_under", "first_half_total", "h1_over_under"],
         _OVER_UNDER),
    ],

    # ── Tennis ────────────────────────────────────────────────────────────────
    "tennis": [
        ("Match Winner",
         ["match_winner", "moneyline", "winner", "tennis_winner",
          "outright_winner", "1x2"],
         _WIN_LOSE),
        ("Set Betting",
         ["set_betting", "correct_score_sets", "sets_score"],
         [("2-0", ["2-0", "2_0"]),
          ("2-1", ["2-1", "2_1"]),
          ("0-2", ["0-2", "0_2"]),
          ("1-2", ["1-2", "1_2"])]),
        ("Total Games Over/Under",
         ["total_games", "over_under_games", "games_over_under"],
         _OVER_UNDER),
        ("Game Handicap",
         ["game_handicap", "games_handicap", "tennis_handicap"],
         [("Home", ["home", "1", "player1"]),
          ("Away", ["away", "2", "player2"])]),
    ],

    # ── Cricket ───────────────────────────────────────────────────────────────
    "cricket": [
        ("Match Winner",
         ["match_winner", "winner", "1x2", "cricket_winner",
          "moneyline", "outright"],
         _WIN_DRAW_WIN),
        ("Match Winner (no draw)",
         ["match_winner_nodraw", "to_win_match", "win_toss_win_match"],
         _WIN_LOSE),
        ("Over/Under Runs",
         ["over_under_runs", "total_runs", "runs_over_under",
          "over_under", "innings_runs"],
         _OVER_UNDER),
        ("Top Batsman",
         ["top_batsman", "top_run_scorer"],
         []),
    ],

    # ── Ice Hockey ────────────────────────────────────────────────────────────
    "ice-hockey": [
        ("Full-Time 1X2",
         ["1x2", "match_winner", "moneyline", "puck_line"],
         _WIN_DRAW_WIN),
        ("Both Teams to Score",
         ["btts", "both_teams_to_score"],
         _YES_NO),
        ("Over/Under Goals",
         ["over_under", "over_under_goals", "over_under_goals_5_5",
          "over_under_5_5", "total_goals"],
         _OVER_UNDER),
        ("Asian Handicap (Puck Line)",
         ["asian_handicap", "puck_line", "handicap"],
         _WIN_LOSE),
        ("60 Min Result",
         ["60_min_result", "regulation_result", "3_way"],
         _WIN_DRAW_WIN),
    ],

    # ── Volleyball ────────────────────────────────────────────────────────────
    "volleyball": [
        ("Match Winner",
         ["match_winner", "winner", "volleyball_winner",
          "moneyline", "1x2"],
         _WIN_LOSE),
        ("Set Handicap",
         ["set_handicap", "handicap_sets", "volleyball_handicap"],
         [("Home -1.5", ["home", "1", "-1.5"]),
          ("Away +1.5", ["away", "2", "+1.5"])]),
        ("Total Sets",
         ["total_sets", "over_under_sets", "sets_over_under"],
         _OVER_UNDER),
        ("Correct Score (Sets)",
         ["correct_score_sets", "sets_score"],
         [("3-0", ["3-0", "3_0"]),
          ("3-1", ["3-1", "3_1"]),
          ("3-2", ["3-2", "3_2"]),
          ("0-3", ["0-3", "0_3"]),
          ("1-3", ["1-3", "1_3"]),
          ("2-3", ["2-3", "2_3"])]),
    ],

    # ── Rugby ─────────────────────────────────────────────────────────────────
    "rugby": [
        ("Match Winner",
         ["match_winner", "moneyline", "1x2", "rugby_winner"],
         _WIN_DRAW_WIN),
        ("Handicap",
         ["asian_handicap", "handicap", "rugby_handicap"],
         _WIN_LOSE),
        ("Over/Under Points",
         ["over_under", "total_points", "over_under_points"],
         _OVER_UNDER),
        ("Both Teams to Score",
         ["btts", "both_teams_to_score"],
         _YES_NO),
    ],

    # ── Handball ──────────────────────────────────────────────────────────────
    "handball": [
        ("Match Winner",
         ["match_winner", "1x2", "moneyline", "handball_winner"],
         _WIN_DRAW_WIN),
        ("Asian Handicap",
         ["asian_handicap", "handicap"],
         _WIN_LOSE),
        ("Over/Under Goals",
         ["over_under", "over_under_goals", "total_goals"],
         _OVER_UNDER),
        ("Both Teams to Score",
         ["btts", "both_teams_to_score"],
         _YES_NO),
        ("Half-Time",
         ["half_time", "half_time_result"],
         _WIN_DRAW_WIN),
    ],

    # ── Table Tennis ─────────────────────────────────────────────────────────
    "table-tennis": [
        ("Match Winner",
         ["match_winner", "winner", "moneyline", "1x2",
          "table_tennis_winner"],
         _WIN_LOSE),
        ("Total Points",
         ["total_points", "over_under", "over_under_points"],
         _OVER_UNDER),
        ("Handicap (Games)",
         ["asian_handicap", "handicap"],
         _WIN_LOSE),
    ],

    # ── MMA ───────────────────────────────────────────────────────────────────
    "mma": [
        ("Fight Winner",
         ["match_winner", "fight_winner", "winner", "moneyline",
          "mma_winner", "1x2"],
         _WIN_LOSE),
        ("Over/Under Rounds",
         ["over_under", "total_rounds", "over_under_rounds"],
         _OVER_UNDER),
        ("Method of Victory",
         ["method_of_victory", "winning_method"],
         [("KO/TKO", ["ko", "tko", "ko_tko"]),
          ("Submission", ["submission", "sub"]),
          ("Decision", ["decision", "dec", "points"])]),
    ],

    # ── Boxing ────────────────────────────────────────────────────────────────
    "boxing": [
        ("Fight Winner",
         ["match_winner", "fight_winner", "winner", "moneyline",
          "boxing_winner", "1x2"],
         _WIN_DRAW_WIN),
        ("Over/Under Rounds",
         ["over_under", "total_rounds", "over_under_rounds"],
         _OVER_UNDER),
        ("Method of Victory",
         ["method_of_victory", "winning_method"],
         [("KO/TKO", ["ko", "tko", "ko_tko"]),
          ("Points", ["decision", "dec", "points"]),
          ("Draw", ["draw", "x"])]),
    ],

    # ── Darts ─────────────────────────────────────────────────────────────────
    "darts": [
        ("Match Winner",
         ["match_winner", "winner", "moneyline", "1x2", "darts_winner"],
         _WIN_LOSE),
        ("Correct Score (Legs)",
         ["correct_score", "legs_score", "correct_score_legs"],
         [("Home", ["home", "1"]),
          ("Away", ["away", "2"])]),
        ("Total 180s",
         ["total_180s", "over_under_180s"],
         _OVER_UNDER),
    ],

    # ── American Football ─────────────────────────────────────────────────────
    "american-football": [
        ("Moneyline",
         ["moneyline", "match_winner", "1x2", "nfl_winner"],
         _WIN_LOSE),
        ("Point Spread",
         ["point_spread", "asian_handicap", "handicap"],
         [("Home", ["home", "1", "h"]),
          ("Away", ["away", "2", "a"])]),
        ("Over/Under Total Points",
         ["over_under", "total_points", "over_under_points"],
         _OVER_UNDER),
        ("1st Half Moneyline",
         ["first_half_1x2", "1h_moneyline", "first_half_winner"],
         _WIN_LOSE),
    ],

    # ── Baseball ─────────────────────────────────────────────────────────────
    "baseball": [
        ("Moneyline",
         ["moneyline", "match_winner", "1x2", "run_line"],
         _WIN_LOSE),
        ("Run Line",
         ["run_line", "point_spread", "asian_handicap"],
         [("Home -1.5", ["home", "1"]),
          ("Away +1.5", ["away", "2"])]),
        ("Over/Under Runs",
         ["over_under", "total_runs", "over_under_runs"],
         _OVER_UNDER),
    ],

    # ── eSoccer ── (same as soccer)
    "esoccer": [],  # filled dynamically below
}

# eSoccer = soccer markets
_SPORT_MARKET_DEFS["esoccer"] = list(_SPORT_MARKET_DEFS["soccer"])

# Kept for backward compatibility (used elsewhere)
BASE_MARKETS = _SPORT_MARKET_DEFS["soccer"]


def _get_display_markets(
    sport:         str,
    matches:       list,
    market_filter: list | None = None,
) -> list:
    """
    Build the ordered list of (label, aliases, outcomes) to display in the
    Word document for a given sport.

    Strategy (3 layers):
      1. Start with hardcoded sport-specific definitions (most important markets
         shown first in the correct order).
      2. Detect any additional market slugs present in the actual match data
         that aren't covered by layer 1 (dynamic — handles harvester variations).
      3. If market_filter is provided, restrict to matching markets only.
    """
    # Layer 1 — sport-specific hardcoded definitions
    base = list(_SPORT_MARKET_DEFS.get(sport.lower(), _SPORT_MARKET_DEFS["soccer"]))

    # Track all alias keys already covered so we don't double-render
    covered: set[str] = set()
    for _, aliases, _ in base:
        covered.update(aliases)

    # Layer 2 — dynamic detection from actual match data
    # Scan `best` to find market slugs not yet covered, then infer outcomes
    dynamic_markets: dict[str, set[str]] = {}  # slug → {outcome_key, …}
    for m in matches:
        for mkt_slug, outcomes_dict in (m.get("best") or {}).items():
            if not isinstance(outcomes_dict, dict):
                continue
            # Skip if already covered by a hardcoded alias
            if mkt_slug in covered:
                continue
            # Also skip if it matches any covered alias (partial)
            already = False
            for cv in covered:
                if cv in mkt_slug or mkt_slug in cv:
                    already = True; break
            if already:
                continue
            dynamic_markets.setdefault(mkt_slug, set()).update(outcomes_dict.keys())

    for mkt_slug, outcome_keys in sorted(dynamic_markets.items()):
        if not outcome_keys:
            continue
        label    = mkt_slug.replace("_", " ").title()
        outcomes = [(k.title(), [k, k.lower()]) for k in sorted(outcome_keys)]
        base.append((label, [mkt_slug], outcomes))
        covered.add(mkt_slug)

    # Layer 3 — apply market_filter if provided
    if not market_filter:
        return base

    mf_set = {m.lower() for m in market_filter}
    return [
        d for d in base
        if any(alias in mf_set for alias in d[1])
        or d[0].lower().replace(" ", "_") in mf_set
    ]


# ── Helpers ────────────────────────────────────────────────────────────────────

def _now_utc() -> datetime:
    return datetime.now(timezone.utc)


def _parse_dt(s: str) -> Optional[datetime]:
    if not s:
        return None
    try:
        dt = datetime.fromisoformat(str(s).strip().replace("Z", "+00:00"))
        return dt.replace(tzinfo=timezone.utc) if dt.tzinfo is None else dt
    except Exception:
        return None


def _eat(dt: datetime) -> datetime:
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt + EAT


def _group_for_dt(dt: datetime) -> Optional[str]:
    h = _eat(dt).hour
    for gid, (_, s, e) in TIME_GROUPS.items():
        if s <= h < e:
            return gid
    return None


def _name_to_slug(name: str) -> str:
    """Convert bookmaker display name to slug."""
    key = name.lower().strip().replace(" ", "").replace("-", "")
    return _BK_NAME_TO_SLUG.get(key, name.lower().replace(" ", "_")[:12])


def _extract_price(val) -> float:
    """Robustly extract a float price from various stored formats."""
    if val is None:
        return 0.0
    if isinstance(val, (int, float)):
        return float(val)
    if isinstance(val, dict):
        for k in ("price", "odd", "odds", "value", "best_price"):
            if val.get(k):
                try:
                    return float(val[k])
                except (TypeError, ValueError):
                    pass
    try:
        return float(val)
    except (TypeError, ValueError):
        return 0.0


def _flatten_markets_json(markets_json) -> dict:
    """
    Normalize BookmakerMatchOdds.markets_json into:
        {market_slug: {outcome_key: float_price}}
    Handles nested formats produced by different harvesters.
    """
    if not markets_json or not isinstance(markets_json, dict):
        return {}
    result: dict = {}
    for mkt_key, mkt_val in markets_json.items():
        if not mkt_key:
            continue
        if isinstance(mkt_val, dict):
            outcomes: dict = {}
            for out_key, out_val in mkt_val.items():
                if out_key in ("name", "id", "status", "specifier"):
                    continue
                price = _extract_price(out_val)
                if price > 1.0:
                    outcomes[str(out_key)] = price
            if outcomes:
                result[str(mkt_key)] = outcomes
    return result


def _extract_odd(
    match: dict,
    bk_slug: str,
    mkt_aliases: list,
    out_aliases: list,
) -> Optional[float]:
    """Pull a single odds value for (bookmaker, market, outcome)."""
    out_set = {a.lower() for a in out_aliases}

    # 1. bookmakers → markets
    bk_data  = (match.get("bookmakers") or {}).get(bk_slug, {})
    mkt_data = bk_data.get("markets") or {} if isinstance(bk_data, dict) else {}

    for key in mkt_aliases:
        mkt = mkt_data.get(key)
        if not isinstance(mkt, dict):
            continue
        for k, v in mkt.items():
            if str(k).lower() not in out_set:
                continue
            fv = _extract_price(v)
            if fv > 1.0:
                return round(fv, 2)

    # 2. markets_by_bk fallback
    mbk = (match.get("markets_by_bk") or {}).get(bk_slug, {})
    for key in mkt_aliases:
        mkt = mbk.get(key)
        if not isinstance(mkt, dict):
            continue
        for k, v in mkt.items():
            if str(k).lower() not in out_set:
                continue
            fv = _extract_price(v)
            if fv > 1.0:
                return round(fv, 2)

    # 3. best odds — only if bk matches
    best = match.get("best") or {}
    for key in mkt_aliases:
        mkt_best = best.get(key) or {}
        for out_label, out_keys in [(a, [a]) for a in out_aliases]:
            for ok in [out_label.lower()] + [a.lower() for a in out_aliases]:
                entry = mkt_best.get(ok)
                if isinstance(entry, dict) and entry.get("bk", "").lower() == bk_slug.lower():
                    fv = _extract_price(entry.get("odd", 0))
                    if fv > 1.0:
                        return round(fv, 2)

    return None


def _detect_active_bks(matches: list) -> list:
    """Return ordered list of bookmaker slugs present in this match list."""
    seen: dict[str, int] = {}
    for m in matches:
        for slug in (m.get("bookmakers") or {}):
            seen[slug] = seen.get(slug, 0) + 1
    result = [b for b in PRIMARY_BKS if b in seen]
    extras = sorted([b for b in seen if b not in PRIMARY_BKS], key=lambda x: -seen[x])
    return result + extras[:3]


def _get_game_ids(m: dict) -> dict:
    """
    Extract short bookmaker game IDs from a match dict.
    Returns {slug: id_str} for slugs with a valid 1-8 digit numeric ID.
    """
    ids: dict = {}

    for slug, val in (m.get("bk_ids") or {}).items():
        if val and str(val).strip().isdigit():
            s = str(val).strip()
            if 1 <= len(s) <= 8:
                ids[slug] = s

    sp_cands = [m.get("sp_game_id"), m.get("sms_id")]
    bt_cands = [m.get("bt_game_id"), m.get("bt_match_id"), m.get("bt_parent_id")]
    od_cands = [m.get("od_game_id"), m.get("od_event_id"), m.get("od_match_id")]

    for slug, cands in [("sp", sp_cands), ("bt", bt_cands), ("od", od_cands)]:
        if slug not in ids:
            for v in cands:
                if v and str(v).strip().isdigit():
                    s = str(v).strip()
                    if 1 <= len(s) <= 8:
                        ids[slug] = s
                        break
    return ids


# ══════════════════════════════════════════════════════════════════════════════
# DATA LOADING — DB primary, Redis fallback
# ══════════════════════════════════════════════════════════════════════════════

def _load_matches(sport: str, date_str: str) -> list:
    """
    Load matches for a sport + EAT date.
    Tries DB first (works for all sports), falls back to Redis.
    """
    matches = _load_from_db(sport, date_str)
    if matches:
        return matches
    return _load_from_redis(sport, date_str)


def _load_from_db(sport: str, date_str: str) -> list:
    """
    Load matches from PostgreSQL.
    Reads UnifiedMatch + BookmakerMatchOdds for the given EAT date window.
    """
    try:
        from app.models.odds import UnifiedMatch, BookmakerMatchOdds
        from app.models.bookmakers_model import Bookmaker
        from app.extensions import db
        from sqlalchemy import or_

        # Parse EAT date → UTC window
        try:
            eat_date = datetime.strptime(date_str, "%Y-%m-%d")
        except Exception:
            eat_date = (_now_utc() + EAT).replace(
                hour=0, minute=0, second=0, microsecond=0
            )

        utc_start = eat_date.replace(tzinfo=timezone.utc) - EAT
        utc_end   = utc_start + timedelta(hours=24)

        # Sport name filter
        sport_names   = _SPORT_DB_NAMES.get(sport.lower(), [sport.replace("-", " ").title()])
        sport_filters = [UnifiedMatch.sport_name.ilike(f"%{n}%") for n in sport_names]

        # Query upcoming + live matches in the window
        ums = (
            UnifiedMatch.query
            .filter(
                UnifiedMatch.start_time >= utc_start,
                UnifiedMatch.start_time < utc_end,
                or_(*sport_filters),
                ~UnifiedMatch.status.in_(["finished", "ft", "complete", "ended", "FT"]),
            )
            .order_by(UnifiedMatch.start_time)
            .all()
        )

        if not ums:
            return []

        # Bulk-load bookmaker odds
        um_ids   = [um.id for um in ums]
        bmo_list = BookmakerMatchOdds.query.filter(
            BookmakerMatchOdds.match_id.in_(um_ids)
        ).all()

        # Bookmaker id → object
        bk_ids  = {bmo.bookmaker_id for bmo in bmo_list}
        bk_map  = {b.id: b for b in Bookmaker.query.filter(Bookmaker.id.in_(bk_ids)).all()} if bk_ids else {}

        # Group BMOs by match id
        bmos_by_match: dict[int, list] = {}
        for bmo in bmo_list:
            bmos_by_match.setdefault(bmo.match_id, []).append(bmo)

        result: list[dict] = []
        for um in ums:
            bookmakers: dict = {}
            best: dict = {}

            for bmo in bmos_by_match.get(um.id, []):
                bk_obj = bk_map.get(bmo.bookmaker_id)
                if not bk_obj:
                    continue
                slug    = _name_to_slug(bk_obj.name)
                markets = _flatten_markets_json(bmo.markets_json or {})
                if not markets:
                    # Try the bmo.markets field if markets_json is empty
                    try:
                        raw_m = getattr(bmo, "markets", None)
                        if isinstance(raw_m, dict):
                            markets = _flatten_markets_json(raw_m)
                    except Exception:
                        pass
                if not markets:
                    continue

                if slug in bookmakers:
                    bookmakers[slug]["markets"].update(markets)
                else:
                    bookmakers[slug] = {
                        "bookmaker": bk_obj.name,
                        "slug":      slug,
                        "markets":   markets,
                    }

                # Build best odds
                for mkt, outcomes in markets.items():
                    best.setdefault(mkt, {})
                    for out, price_data in outcomes.items():
                        pf = _extract_price(price_data)
                        if pf > 1.0:
                            existing = best[mkt].get(out)
                            if not existing or pf > existing.get("odd", 0):
                                best[mkt][out] = {"odd": pf, "bk": slug}

            # Supplement with aggregated markets_json on UnifiedMatch itself
            try:
                agg_json = getattr(um, "markets_json", None)
                if agg_json:
                    agg = _flatten_markets_json(agg_json)
                    for mkt, outcomes in agg.items():
                        best.setdefault(mkt, {})
                        for out, price in outcomes.items():
                            pf = _extract_price(price)
                            if pf > 1.0 and out not in best.get(mkt, {}):
                                best[mkt][out] = {"odd": pf, "bk": "agg"}
            except Exception:
                pass

            status = (getattr(um, "status", None) or "upcoming").lower()
            result.append({
                "match_id":        str(um.id),
                "join_key":        f"br_{um.parent_match_id}" if um.parent_match_id else f"db_{um.id}",
                "parent_match_id": um.parent_match_id or str(um.id),
                "betradar_id":     um.parent_match_id or "",
                "home_team":       um.home_team_name or "",
                "away_team":       um.away_team_name or "",
                "competition":     um.competition_name or "",
                "sport":           sport,
                "sport_name":      getattr(um, "sport_name", sport) or sport,
                "start_time":      um.start_time.isoformat() if um.start_time else "",
                "status":          status,
                "is_live":         status in ("in_play", "live", "in_progress"),
                "bookmakers":      bookmakers,
                "best":            best,
                "bk_count":        len(bookmakers),
                "market_slugs":    list(best.keys()),
                "has_arb":         False,
                "arb_opportunities": [],
            })

        return result

    except Exception:
        import traceback
        traceback.print_exc()
        return []


def _load_from_redis(sport: str, date_str: str) -> list:
    """Redis fallback — existing unified cache approach."""
    try:
        from datetime import datetime as _dt
        eat_date = _dt.strptime(date_str, "%Y-%m-%d")
        utc_start = eat_date.replace(tzinfo=timezone.utc) - EAT
        utc_end   = utc_start + timedelta(hours=24)

        matches_raw: list = []
        try:
            from app.api.odds_stream import _get_unified_patched
            for mode in ("upcoming", "live"):
                for m in _get_unified_patched(mode, sport, force_refresh=False):
                    if not isinstance(m, dict):
                        continue
                    dt = _parse_dt(m.get("start_time", ""))
                    if dt and utc_start <= dt < utc_end:
                        matches_raw.append(m)
        except Exception:
            pass

        # Deduplicate
        seen: set = set()
        out:  list = []
        for m in matches_raw:
            jk = m.get("join_key") or m.get("parent_match_id") or m.get("match_id")
            if jk and jk in seen:
                continue
            if jk:
                seen.add(jk)
            out.append(m)

        out.sort(key=lambda x: x.get("start_time") or "")
        return out

    except Exception:
        return []


# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC API
# ══════════════════════════════════════════════════════════════════════════════

def get_available_groups(sport: str, date_str: str) -> list:
    """
    Returns a list of group dicts for the given sport + date.
    Only groups with ≥1 match are returned.
    [{id, label, count, time_range, earliest, latest}, …]
    """
    matches = _load_matches(sport, date_str)
    buckets: dict[str, list] = {gid: [] for gid in TIME_GROUPS}

    for m in matches:
        dt = _parse_dt(m.get("start_time", ""))
        if not dt:
            continue
        gid = _group_for_dt(dt)
        if gid:
            buckets[gid].append(m)

    result = []
    for gid, (label, start_h, end_h) in TIME_GROUPS.items():
        ms = buckets[gid]
        if not ms:
            continue
        dts = [_parse_dt(m.get("start_time", "")) for m in ms]
        dts = [d for d in dts if d]
        result.append({
            "id":         gid,
            "label":      label,
            "count":      len(ms),
            "time_range": f"{start_h:02d}:00–{end_h:02d}:00 EAT",
            "earliest":   _eat(min(dts)).strftime("%H:%M") if dts else "",
            "latest":     _eat(max(dts)).strftime("%H:%M") if dts else "",
        })
    return result


def generate_group_document(
    sport:          str,
    group_id:       str,
    date_str:       str,
    market_filter:  Optional[list] = None,
) -> io.BytesIO:
    """
    Generate a single Word document for one time group.
    Styled like the Soccer Betting Analysis image:
      - Dark theme, portrait A4, ultra-compact
      - Per-match tables: Market | Selection | SP | BT | OD …
      - Best-odd cell highlighted green
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

    # ── Load data ─────────────────────────────────────────────────────────────
    all_matches = _load_matches(sport, date_str)
    group_label, grp_start_h, grp_end_h = TIME_GROUPS.get(group_id, ("All Matches", 0, 24))

    group_matches = [
        m for m in all_matches
        if (dt := _parse_dt(m.get("start_time", ""))) is not None
        and _group_for_dt(dt) == group_id
    ]
    group_matches.sort(key=lambda x: x.get("start_time") or "")

    active_bks = _detect_active_bks(group_matches) or PRIMARY_BKS[:]

    # ── Sport-aware market definitions (replaces static BASE_MARKETS) ─────────
    # _get_display_markets uses sport-specific hardcoded definitions for the
    # most important markets, then dynamically adds any extra markets it finds
    # in the actual match data.  This means basketball gets match_winner /
    # total_points, cricket gets match_winner / runs, etc.
    display_markets = _get_display_markets(sport, group_matches, market_filter)

    # ── Style constants ───────────────────────────────────────────────────────
    FF     = "Arial"
    C_HDR  = "0F172A"
    C_ROW0 = "0D1B2E"
    C_ROW1 = "1E293B"
    C_THDR = "1B2A4A"
    C_BEST = "14532D"
    C_LINE = "334155"

    W    = RGBColor(0xFF, 0xFF, 0xFF)
    CYAN = RGBColor(0x38, 0xBD, 0xF8)
    GOLD = RGBColor(0xFB, 0xBF, 0x24)
    MUTED= RGBColor(0x64, 0x74, 0x8B)
    GREEN= RGBColor(0x4A, 0xDE, 0x80)
    DIM  = RGBColor(0x94, 0xA3, 0xB8)

    SPORT_EMOJI = {
        "soccer": "⚽", "basketball": "🏀", "tennis": "🎾",
        "ice-hockey": "🏒", "volleyball": "🏐", "cricket": "🏏",
        "rugby": "🏉", "table-tennis": "🏓", "handball": "🤾",
        "boxing": "🥊", "mma": "🥊", "darts": "🎯",
        "american-football": "🏈", "baseball": "⚾",
    }
    sport_emoji = SPORT_EMOJI.get(sport.lower(), "🏆")

    # ── XML helpers ───────────────────────────────────────────────────────────
    def _shd(cell, hex_col):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_col}"/>'))

    def _pad(cell, top=45, bot=45, left=60, right=60):
        tcPr  = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for nm, v in [("top", top), ("bottom", bot), ("left", left), ("right", right)]:
            n = OxmlElement(f"w:{nm}")
            n.set(qn("w:w"), str(v)); n.set(qn("w:type"), "dxa")
            tcMar.append(n)
        tcPr.append(tcMar)

    def _borders(cell, color=C_LINE, sz="2"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for nm in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{nm}")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "0"); b.set(qn("w:color"), color)
            tcB.append(b)
        tcPr.append(tcB)

    def _no_borders(table):
        tblPr = table._tbl.tblPr
        tcB   = OxmlElement("w:tblBorders")
        for nm in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{nm}"); b.set(qn("w:val"), "none")
            tcB.append(b)
        tblPr.append(tcB)

    def _cw(cell, text, bold=False, italic=False,
             color=None, size=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             sa=0, sb=0):
        p = (cell.paragraphs[0]
             if (cell.paragraphs and not cell.paragraphs[0].text)
             else cell.add_paragraph())
        p.alignment = align
        p.paragraph_format.space_after  = Pt(sa)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.name = FF
        if color: r.font.color.rgb = color
        if size:  r.font.size = size
        return p

    # ── Page setup: Portrait A4, ultra-tight margins ──────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = Cm(0.55)
        sec.bottom_margin = Cm(0.50)
        sec.left_margin   = Cm(0.65)
        sec.right_margin  = Cm(0.65)

    doc.styles["Normal"].font.name      = FF
    doc.styles["Normal"].font.size      = Pt(6.5)
    doc.styles["Normal"].font.color.rgb = W

    # ── 2-column layout ───────────────────────────────────────────────────────
    # Header (full-width) renders in section 1 (1 col).
    # A continuous section break after the header switches to 2 columns.
    # Per-column usable width: (19.7 - 0.5 gutter) / 2 = 9.6 cm
    COL_W_CM   = 9.6          # usable width per column
    USABLE_W   = Cm(19.7)     # used for the full-width header table only
    COL_USABLE = Cm(COL_W_CM) # used for per-match tables

    n_bk  = len(active_bks)
    # Column widths within each 9.6 cm match table
    mkt_w = 3.3
    sel_w = 1.3
    bk_w  = round((COL_W_CM - mkt_w - sel_w) / max(n_bk, 1), 2)
    bk_w  = max(bk_w, 1.5)
    mkt_w = max(round(COL_W_CM - sel_w - bk_w * n_bk, 2), 2.5)
    n_cols = 2 + n_bk

    # ── Header ────────────────────────────────────────────────────────────────
    now_eat = _eat(_now_utc())
    try:
        eat_date_display = datetime.strptime(date_str, "%Y-%m-%d").strftime("%A, %d %B %Y")
    except Exception:
        eat_date_display = date_str

    ht = doc.add_table(rows=1, cols=3)
    _no_borders(ht); ht.autofit = False
    ht.columns[0].width = Cm(9.5)
    ht.columns[1].width = Cm(3.5)
    ht.columns[2].width = Cm(6.7)

    hl, hm, hr = ht.rows[0].cells[0], ht.rows[0].cells[1], ht.rows[0].cells[2]
    for c in (hl, hm, hr):
        _shd(c, C_HDR); _pad(c, top=120, bot=100, left=130, right=100)

    # Left: title
    p_title = hl.paragraphs[0]
    p_title.paragraph_format.line_spacing = 1.0
    p_title.paragraph_format.space_after  = Pt(3)
    rt1 = p_title.add_run(f"{sport_emoji} {sport.upper()} BETTING ANALYSIS")
    rt1.bold = True; rt1.font.size = Pt(13); rt1.font.color.rgb = W; rt1.font.name = FF
    p_sub = hl.add_paragraph()
    p_sub.paragraph_format.line_spacing = 1.0; p_sub.paragraph_format.space_after = Pt(0)
    rs1 = p_sub.add_run(f"📋 {group_label}   |   📅 {eat_date_display}")
    rs1.font.size = Pt(6.5); rs1.font.color.rgb = MUTED; rs1.font.name = FF

    # Middle: match count
    _shd(hm, C_HDR); _pad(hm, top=110, bot=100, left=80, right=80)
    pc = hm.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.line_spacing = 1.0
    rc1 = pc.add_run(str(len(group_matches)))
    rc1.bold = True; rc1.font.size = Pt(22); rc1.font.color.rgb = CYAN; rc1.font.name = FF
    pc2 = hm.add_paragraph()
    pc2.alignment = WD_ALIGN_PARAGRAPH.CENTER; pc2.paragraph_format.line_spacing = 1.0
    rc2 = pc2.add_run("MATCHES"); rc2.font.size = Pt(6); rc2.font.color.rgb = MUTED; rc2.font.name = FF

    # Right: bookmaker legend
    _pad(hr, top=110, bot=100, left=80, right=120)
    pl = hr.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT; pl.paragraph_format.line_spacing = 1.0
    pl_title = pl.add_run("Bookmakers:  ")
    pl_title.font.size = Pt(6.5); pl_title.font.color.rgb = MUTED; pl_title.font.name = FF
    for slug in active_bks:
        hx, _ = BK_COLORS.get(slug, ("475569", "F1F5F9"))
        lbl   = BK_LABELS.get(slug, slug.upper()[:3])
        rb = pl.add_run(f"[{lbl}]  ")
        rb.bold = True; rb.font.size = Pt(7)
        rb.font.color.rgb = RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))
        rb.font.name = FF
    pl2 = hr.add_paragraph()
    pl2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; pl2.paragraph_format.line_spacing = 1.0
    rl2 = pl2.add_run(f"🟢 Green = best odd   |   Generated {now_eat.strftime('%H:%M')} EAT")
    rl2.font.size = Pt(5.5); rl2.font.color.rgb = MUTED; rl2.font.name = FF

    # Accent stripe
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(0); pa.paragraph_format.space_after = Pt(0)
    ra = pa.add_run("─" * 200)
    ra.font.size = Pt(1.5); ra.font.color.rgb = CYAN; ra.font.name = FF

    # ── Switch to 2-column layout for match content ───────────────────────────
    # Insert a continuous section break. The sectPr inside the paragraph's pPr
    # describes the section that ENDS here (1-col header section). Everything
    # after this paragraph is section 2 which has 2 columns.
    def _insert_two_col_break():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr    = p._p.get_or_add_pPr()
        sectPr = OxmlElement("w:sectPr")
        # 2-column spec for section 2
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "283")   # ≈ 0.5 cm gutter
        cols.set(qn("w:equalWidth"), "1")
        sectPr.append(cols)
        # Continuous break — no page break between header and match content
        t = OxmlElement("w:type")
        t.set(qn("w:val"), "continuous")
        sectPr.append(t)
        pPr.append(sectPr)

    if not group_matches:
        pn = doc.add_paragraph()
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = pn.add_run("No matches in this time slot.")
        rn.italic = True; rn.font.color.rgb = MUTED; rn.font.name = FF
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    _insert_two_col_break()

    # ── Match tables ──────────────────────────────────────────────────────────
    for match_idx, m in enumerate(group_matches, 1):
        m_dt  = _parse_dt(m.get("start_time", "")) or _now_utc()
        home  = (m.get("home_team") or "Home")[:24]
        away  = (m.get("away_team") or "Away")[:24]
        comp  = (m.get("competition") or "")[:36]
        ko    = _eat(m_dt).strftime("%H:%M")
        has_arb = bool(m.get("has_arb") or m.get("arbitrage"))
        game_ids = _get_game_ids(m)

        # Match header
        mh_bg = "064E3B" if has_arb else C_HDR
        mht = doc.add_table(rows=1, cols=1)
        _no_borders(mht); mht.autofit = False
        mht.columns[0].width = COL_USABLE
        mhc = mht.rows[0].cells[0]
        _shd(mhc, mh_bg)
        _pad(mhc, top=70, bot=55, left=90, right=80)
        _borders(mhc, "1E40AF" if not has_arb else "16A34A", "6")

        pm = mhc.paragraphs[0]
        pm.paragraph_format.line_spacing = 1.0; pm.paragraph_format.space_after = Pt(1)
        rnum = pm.add_run(f"{match_idx}. ")
        rnum.font.size = Pt(6.5); rnum.font.color.rgb = MUTED; rnum.font.name = FF
        rh = pm.add_run(home); rh.bold = True; rh.font.size = Pt(8.5)
        rh.font.color.rgb = CYAN; rh.font.name = FF
        rv = pm.add_run("  vs  "); rv.font.size = Pt(6); rv.font.color.rgb = MUTED; rv.font.name = FF
        ra_ = pm.add_run(away); ra_.bold = True; ra_.font.size = Pt(8.5)
        ra_.font.color.rgb = W; ra_.font.name = FF
        rko = pm.add_run(f"   ⏱ {ko}")
        rko.font.size = Pt(6); rko.font.color.rgb = GOLD; rko.font.name = FF
        if comp:
            rcp = pm.add_run(f"   🏆 {comp}")
            rcp.font.size = Pt(5.5); rcp.font.color.rgb = DIM; rcp.font.name = FF
        if has_arb:
            rarb = pm.add_run("  ⚡ARB")
            rarb.bold = True; rarb.font.size = Pt(6)
            rarb.font.color.rgb = GREEN; rarb.font.name = FF

        # Game IDs row (only shown when at least one BK ID is known)
        if game_ids:
            p_ids = mhc.add_paragraph()
            p_ids.paragraph_format.line_spacing = 1.0
            p_ids.paragraph_format.space_after  = Pt(0)
            id_parts = []
            for slug in active_bks:
                if slug in game_ids:
                    id_parts.append(f"{slug.upper()}#{game_ids[slug]}")
            if id_parts:
                r_ids = p_ids.add_run("📲  " + "   ".join(id_parts))
                r_ids.font.size = Pt(5.5)
                r_ids.font.color.rgb = RGBColor(0xA7, 0xF3, 0xD0)
                r_ids.font.name = FF

        # Market table
        tbl = doc.add_table(rows=1, cols=n_cols)
        _no_borders(tbl); tbl.autofit = False
        tbl.columns[0].width = Cm(mkt_w)
        tbl.columns[1].width = Cm(sel_w)
        for ci in range(2, n_cols):
            tbl.columns[ci].width = Cm(bk_w)

        # Column headers
        hrow = tbl.rows[0]
        _shd(hrow.cells[0], C_THDR); _pad(hrow.cells[0], 50, 50, 60, 45); _borders(hrow.cells[0], C_LINE, "2")
        _cw(hrow.cells[0], "Market", bold=True, color=CYAN, size=Pt(6))
        _shd(hrow.cells[1], C_THDR); _pad(hrow.cells[1], 50, 50, 45, 45); _borders(hrow.cells[1], C_LINE, "2")
        _cw(hrow.cells[1], "Selection", bold=True, color=CYAN, size=Pt(6), align=WD_ALIGN_PARAGRAPH.CENTER)

        for ci, slug in enumerate(active_bks):
            c = hrow.cells[2 + ci]
            hx, _ = BK_COLORS.get(slug, ("334155", "1E293B"))
            _shd(c, hx); _pad(c, 55, 55, 50, 50); _borders(c, C_LINE, "2")
            _cw(c, BK_LABELS.get(slug, slug.upper()[:4]), bold=True, color=W,
                size=Pt(7), align=WD_ALIGN_PARAGRAPH.CENTER)

        # Data rows
        row_idx = 0
        for mkt_label, mkt_aliases, outcomes in display_markets:
            any_data = any(
                _extract_odd(m, slug, mkt_aliases, out_aliases) is not None
                for slug in active_bks
                for _, out_aliases in outcomes
            )
            if not any_data:
                continue

            for out_label, out_aliases in outcomes:
                bk_odds = {slug: _extract_odd(m, slug, mkt_aliases, out_aliases) for slug in active_bks}
                valid_odds = [v for v in bk_odds.values() if v is not None]
                if not valid_odds:
                    continue
                best_odd = max(valid_odds)

                bg = C_ROW0 if row_idx % 2 == 0 else C_ROW1
                dr = tbl.add_row()

                c0 = dr.cells[0]; _shd(c0, bg); _pad(c0, 40, 40, 70, 40); _borders(c0, C_LINE, "2")
                _cw(c0, mkt_label if out_label == outcomes[0][0] else "", color=DIM, size=Pt(6.5))

                c1 = dr.cells[1]; _shd(c1, bg); _pad(c1, 40, 40, 50, 50); _borders(c1, C_LINE, "2")
                _cw(c1, out_label, bold=True, color=W, size=Pt(6.5), align=WD_ALIGN_PARAGRAPH.CENTER)

                for ci, slug in enumerate(active_bks):
                    odd = bk_odds.get(slug)
                    c   = dr.cells[2 + ci]
                    is_best = (odd is not None and odd == best_odd)
                    _shd(c, C_BEST if is_best else bg)
                    _pad(c, 40, 40, 45, 45); _borders(c, C_LINE, "2")
                    if odd is not None:
                        p_odd = c.paragraphs[0]
                        p_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_odd.paragraph_format.line_spacing = 1.0
                        ro = p_odd.add_run(f"{odd:.2f}")
                        ro.bold = is_best
                        ro.font.size = Pt(7 if is_best else 6.5)
                        ro.font.color.rgb = GREEN if is_best else W
                        ro.font.name = FF
                    else:
                        _cw(c, "—", color=MUTED, size=Pt(6), align=WD_ALIGN_PARAGRAPH.CENTER)

                row_idx += 1

        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(5)
        sp.paragraph_format.space_before = Pt(0)

    # ── Footer ────────────────────────────────────────────────────────────────
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.paragraph_format.space_before = Pt(8)
    rf = pf.add_run(
        f"OddsKenya · {sport_emoji} {sport.upper()} · {group_label} · "
        f"{eat_date_display} · Verify all odds before placing bets."
    )
    rf.font.size = Pt(5.5); rf.italic = True; rf.font.color.rgb = MUTED; rf.font.name = FF

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf