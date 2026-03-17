"""
app/workers/bankroll_service.py
================================
Premium bankroll rolling engine.

Features
---------
  • Kelly Criterion stake sizing
  • Flat stake mode
  • Fibonacci progression on loss runs
  • Stop-loss enforcement
  • Recommend best current bets from active EV + Arb opportunities
  • Email notification when bankroll achieves profit milestone
"""

from __future__ import annotations

import math
from datetime import datetime, timezone
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from app.models.bank_roll import BankrollTarget, BankrollAccount


# =============================================================================
# Kelly Criterion
# =============================================================================

def kelly_stake(odds: float, probability: float,
                bankroll: float, fraction: float = 0.5) -> float:
    """
    Return recommended stake using fractional Kelly criterion.

    fraction=0.5 is half-Kelly — recommended for real-money use.
    """
    if odds <= 1.0 or not (0 < probability < 1):
        return 0.0
    b         = odds - 1.0          # net payout per unit
    q         = 1.0 - probability
    kelly_f   = (b * probability - q) / b
    stake_f   = kelly_f * fraction
    return round(max(stake_f, 0) * bankroll, 2)


def fibonacci_stake(base_stake: float, loss_count: int,
                    max_multiplier: float = 8.0) -> float:
    """Fibonacci progression: 1, 1, 2, 3, 5, 8, 13, ... × base_stake (capped)."""
    fib = [1, 1]
    while len(fib) <= loss_count:
        fib.append(fib[-1] + fib[-2])
    mult = min(fib[loss_count], max_multiplier)
    return round(base_stake * mult, 2)


# =============================================================================
# Bankroll Engine
# =============================================================================

class BankrollEngine:
    """
    Recommends optimal bets for a BankrollTarget given current active
    arbitrage and EV opportunities.
    """

    def __init__(self, target: "BankrollTarget"):
        self.target = target

    def recommend(self, max_recs: int = 10) -> list[dict]:
        """
        Return a ranked list of bet recommendations for this target.
        Checks both arb opportunities and +EV bets.
        """
        if self.target.is_stop_loss_triggered:
            return [{
                "warning": "Stop-loss triggered. Bankroll has fallen too far.",
                "stopped_at": datetime.now(timezone.utc).isoformat(),
            }]

        recs: list[dict] = []
        recs.extend(self._arb_recommendations())
        recs.extend(self._ev_recommendations())

        # Sort by expected profit
        recs.sort(key=lambda r: r.get("expected_profit_kes", 0), reverse=True)
        return recs[:max_recs]

    def _total_bankroll(self) -> float:
        """Sum of all user's bankroll accounts."""
        from app.models.bank_roll import BankrollAccount
        accounts = BankrollAccount.query.filter_by(
            user_id=self.target.user_id, is_active=True
        ).all()
        return sum(a.balance_kes for a in accounts)

    def _max_stake(self) -> float:
        total = self._total_bankroll()
        return round(total * self.target.max_bet_pct, 2)

    def _arb_recommendations(self) -> list[dict]:
        from app.models.odds_model import ArbitrageOpportunity, UnifiedMatch

        arbs = ArbitrageOpportunity.query.filter(
            ArbitrageOpportunity.is_active == True,
            ArbitrageOpportunity.max_profit_percentage >= self.target.min_arb_profit,
        ).order_by(ArbitrageOpportunity.max_profit_percentage.desc()).limit(5).all()

        recs: list[dict] = []
        max_stake = self._max_stake()
        total_brl = self._total_bankroll()

        for arb in arbs:
            um = UnifiedMatch.query.get(arb.match_id)
            legs = []
            for leg in arb.legs:
                stake = round(max_stake * (leg.stake_pct or 0) / 100, 2)
                legs.append({
                    "selection":    leg.selection_name,
                    "bookmaker":    leg.bookmaker_name,
                    "odds":         leg.price_at_discovery,
                    "stake_pct":    leg.stake_pct,
                    "stake_kes":    stake,
                })
            recs.append({
                "type":               "arbitrage",
                "market":             arb.market_definition.name if arb.market_definition else "",
                "profit_pct":         arb.max_profit_percentage,
                "total_stake_kes":    max_stake,
                "expected_profit_kes": round(max_stake * arb.max_profit_percentage / 100, 2),
                "legs":               legs,
                "match": {
                    "home":        um.home_team_name if um else "",
                    "away":        um.away_team_name if um else "",
                    "start_time":  um.start_time.isoformat() if um and um.start_time else None,
                },
            })
        return recs

    def _ev_recommendations(self) -> list[dict]:
        from app.models.odds_model import EVOpportunity, UnifiedMatch

        evs = EVOpportunity.query.filter(
            EVOpportunity.is_active == True,
            EVOpportunity.edge_pct  >= self.target.min_edge_pct,
        ).order_by(EVOpportunity.edge_pct.desc()).limit(5).all()

        recs: list[dict] = []
        total_brl = self._total_bankroll()

        for ev in evs:
            stake_kes = kelly_stake(
                odds        = ev.odds,
                probability = ev.no_vig_probability,
                bankroll    = total_brl,
                fraction    = 0.5,
            )
            stake_kes = min(stake_kes, self._max_stake())
            um = UnifiedMatch.query.get(ev.match_id) if ev.match_id else None
            recs.append({
                "type":               "ev",
                "market":             ev.market_definition.name if ev.market_definition else "",
                "selection":          ev.selection,
                "bookmaker":          ev.bookmaker_name,
                "odds":               ev.odds,
                "fair_value":         ev.fair_value_price,
                "edge_pct":           ev.edge_pct,
                "kelly_stake_kes":    stake_kes,
                "expected_profit_kes": round(stake_kes * ev.edge_pct / 100, 2),
                "match": {
                    "home":        um.home_team_name if um else "",
                    "away":        um.away_team_name if um else "",
                    "start_time":  um.start_time.isoformat() if um and um.start_time else None,
                },
            })
        return recs

    def record_bet_result(self, bet_id: int, won: bool, pnl_kes: float) -> None:
        """Update bankroll and record outcome after a bet settles."""
        from app.models.bank_roll import BankrollBet, BankrollTarget
        from app.extensions import db

        bet    = BankrollBet.query.get(bet_id)
        target = self.target

        if not bet:
            return

        bet.status     = "won" if won else "lost"
        bet.pnl_kes    = pnl_kes
        bet.settled_at = datetime.now(timezone.utc)

        target.current_kes += pnl_kes

        if target.current_kes >= target.starting_kes + target.target_kes:
            target.achieved_at = datetime.now(timezone.utc)
            target.is_active   = False

        db.session.commit()


"""
app/workers/notification_service.py
=====================================
Email and push notification dispatcher.
Uses SMTP (configurable) or SendGrid.
"""


class NotificationService:
    """Sends alerts to Pro and Premium users."""

    @staticmethod
    def send_alert(user, match_label: str, arbs: list, evs: list,
                   event_type: str) -> None:
        """Dispatch email alert for arb/EV opportunities."""
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText

        subject, body = NotificationService._build_email(
            user, match_label, arbs, evs, event_type
        )

        smtp_host = __import__("os").environ.get("SMTP_HOST", "smtp.gmail.com")
        smtp_port = int(__import__("os").environ.get("SMTP_PORT", 587))
        smtp_user = __import__("os").environ.get("SMTP_USER", "")
        smtp_pass = __import__("os").environ.get("SMTP_PASS", "")
        from_addr = __import__("os").environ.get("FROM_EMAIL", smtp_user)

        if not smtp_user or not smtp_pass:
            # Silently skip in dev
            return

        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"]    = from_addr
        msg["To"]      = user.email
        msg.attach(MIMEText(body, "html"))

        with smtplib.SMTP(smtp_host, smtp_port) as s:
            s.ehlo()
            s.starttls()
            s.login(smtp_user, smtp_pass)
            s.sendmail(from_addr, [user.email], msg.as_string())

    @staticmethod
    def _build_email(user, match_label: str, arbs: list, evs: list,
                     event_type: str) -> tuple[str, str]:
        name    = user.display_name or user.email.split("@")[0]
        subject = f"⚡ OddsKenya Alert: {event_type.upper()} on {match_label}"

        arb_rows = ""
        for arb in arbs[:3]:
            arb_rows += f"""
            <tr>
              <td>{arb.market_definition.name if arb.market_definition else '—'}</td>
              <td style="color:#22c55e;font-weight:bold">+{arb.max_profit_percentage:.2f}%</td>
              <td>{len(arb.legs)} legs</td>
            </tr>"""

        ev_rows = ""
        for ev in evs[:3]:
            ev_rows += f"""
            <tr>
              <td>{ev.market_definition.name if ev.market_definition else '—'}</td>
              <td>{ev.bookmaker_name}</td>
              <td style="color:#3b82f6;font-weight:bold">+{ev.edge_pct:.2f}%</td>
            </tr>"""

        body = f"""
        <html><body style="font-family:monospace;background:#0a0f0a;color:#e5e5e5;padding:24px">
          <h2 style="color:#F5C842">⚡ OddsKenya Alert</h2>
          <p>Hi {name},</p>
          <p><strong>{match_label}</strong> has a new {event_type} opportunity.</p>

          {"<h3 style='color:#22c55e'>Arbitrage</h3><table border=1 cellpadding=6>" + arb_rows + "</table>" if arb_rows else ""}
          {"<h3 style='color:#3b82f6'>+EV Opportunities</h3><table border=1 cellpadding=6>" + ev_rows + "</table>" if ev_rows else ""}

          <p style="color:#888;font-size:12px">
            Manage notifications at <a href="https://oddskenya.com/settings/notifications">Settings</a>.
          </p>
        </body></html>"""

        return subject, body


"""
app/workers/export_service.py
================================
PDF and DOCX export for premium users.
"""

import io
from flask import Response


class ExportService:
    """Generate downloadable odds reports."""

    @staticmethod
    def generate(matches: list[dict], sport: str, fmt: str) -> Response:
        if fmt == "pdf":
            return ExportService._pdf(matches, sport)
        return ExportService._docx(matches, sport)

    @staticmethod
    def _pdf(matches: list[dict], sport: str) -> Response:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError:
            return Response("reportlab not installed", status=500)

        buf    = io.BytesIO()
        doc    = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story  = []

        now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        story.append(Paragraph(f"OddsKenya — {sport.title()} Upcoming Matches", styles["Title"]))
        story.append(Paragraph(f"Generated: {now_str}", styles["Normal"]))
        story.append(Spacer(1, 12))

        # Table header
        data = [["HOME", "AWAY", "COMPETITION", "START", "H", "D", "A", "O/U"]]
        for m in matches[:200]:
            markets = m.get("markets") or m.get("best_odds") or {}
            ox2     = markets.get("1X2") or markets.get("1x2") or {}
            ou      = next((v for k, v in markets.items() if "over_under" in k), {})

            def _odd(d: dict, key: str) -> str:
                v = d.get(key) or d.get(key.lower())
                if isinstance(v, dict):
                    return f"{v.get('odds') or v.get('odd') or '—'}"
                return str(v) if v else "—"

            data.append([
                (m.get("home_team") or "")[:18],
                (m.get("away_team") or "")[:18],
                (m.get("competition") or "")[:16],
                (m.get("start_time") or "")[:16],
                _odd(ox2, "Home"), _odd(ox2, "Draw"), _odd(ox2, "Away"),
                _odd(ou, "over"),
            ])

        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",  (0, 0), (-1, 0), colors.HexColor("#1a2e1a")),
            ("TEXTCOLOR",   (0, 0), (-1, 0), colors.HexColor("#F5C842")),
            ("FONTSIZE",    (0, 0), (-1, -1), 7),
            ("GRID",        (0, 0), (-1, -1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f0")]),
        ]))
        story.append(t)
        doc.build(story)
        buf.seek(0)

        return Response(
            buf.read(),
            mimetype="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=oddskenya_{sport}_{datetime.now(timezone.utc).strftime('%Y%m%d')}.pdf"},
        )

    @staticmethod
    def _docx(matches: list[dict], sport: str) -> Response:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            return Response("python-docx not installed", status=500)

        doc = Document()
        doc.add_heading(f"OddsKenya — {sport.title()} Upcoming Matches", 0)
        doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, label in enumerate(["HOME", "AWAY", "COMPETITION", "START", "H", "D", "A", "O/U"]):
            hdr[i].text = label

        for m in matches[:200]:
            markets = m.get("markets") or m.get("best_odds") or {}
            ox2     = markets.get("1X2") or markets.get("1x2") or {}
            ou      = next((v for k, v in markets.items() if "over_under" in k), {})

            def _odd(d: dict, key: str) -> str:
                v = d.get(key) or d.get(key.lower())
                if isinstance(v, dict): return str(v.get("odds") or v.get("odd") or "—")
                return str(v) if v else "—"

            row = table.add_row().cells
            row[0].text = (m.get("home_team") or "")[:20]
            row[1].text = (m.get("away_team") or "")[:20]
            row[2].text = (m.get("competition") or "")[:18]
            row[3].text = (m.get("start_time") or "")[:16]
            row[4].text = _odd(ox2, "Home")
            row[5].text = _odd(ox2, "Draw")
            row[6].text = _odd(ox2, "Away")
            row[7].text = _odd(ou, "over")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return Response(
            buf.read(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=oddskenya_{sport}_{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"},
        )